from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / ".docx_build" / "assets"
OUTPUT = ROOT / "CALIBER_STUDIO_MASTER_PLAN.docx"

FONT_REGULAR = Path("/System/Library/Fonts/Supplemental/Arial.ttf")
FONT_BOLD = Path("/System/Library/Fonts/Supplemental/Arial Bold.ttf")

NAVY = "0B1F33"
INK = "172033"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "14B8A6"
MUTED = "667085"
LIGHT = "F2F4F7"
BLUE_GRAY = "E8EEF5"
PALE_TEAL = "E8F8F5"
WHITE = "FFFFFF"
LINE = "D6DCE4"
GOLD = "C58A1B"
RED = "9B1C1C"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    value = hex_value.replace("#", "")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_run_font(
    run,
    name: str = "Calibri",
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
):
    run.font.name = name
    run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def remove_paragraph_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is not None:
        p_pr.remove(p_bdr)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA):
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}, got {sum(widths_dxa)}")

    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[index] / 1440)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    set_table_borders(table)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
    cant_split.set(qn("w:val"), "true")


def add_numbering_definition(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
        if node.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
        if node.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    level.append(num_fmt)

    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    level.append(level_text)

    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.append(r_pr)

    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)


def set_repeat_keep(paragraph, keep_with_next=False, keep_together=False, page_break_before=False):
    paragraph.paragraph_format.keep_with_next = keep_with_next
    paragraph.paragraph_format.keep_together = keep_together
    paragraph.paragraph_format.page_break_before = page_break_before


def add_bullet(doc, text: str, bullet_num_id: int, lead: str | None = None):
    p = doc.add_paragraph(style="Caliber Bullet")
    apply_num(p, bullet_num_id)
    if lead:
        r = p.add_run(lead)
        set_run_font(r, bold=True, color=INK)
        r = p.add_run(text)
        set_run_font(r, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_numbered(doc, text: str, decimal_num_id: int, lead: str | None = None):
    p = doc.add_paragraph(style="Caliber Number")
    apply_num(p, decimal_num_id)
    if lead:
        r = p.add_run(lead)
        set_run_font(r, bold=True, color=INK)
        r = p.add_run(text)
        set_run_font(r, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_body(doc, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph(style="Normal")
    if bold_lead:
        r = p.add_run(bold_lead)
        set_run_font(r, bold=True, color=INK)
        r = p.add_run(text)
        set_run_font(r, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, color=INK)
    return p


def add_callout(doc, label: str, text: str, fill=BLUE_GRAY, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    prevent_row_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:color"), accent)
    borders.append(left)

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(label.upper() + "  ")
    set_run_font(r, size=9.5, bold=True, color=accent)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_table(doc, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header_row = table.rows[0]
    mark_header_row(header_row)
    prevent_row_split(header_row)
    for index, value in enumerate(headers):
        cell = header_row.cells[index]
        set_cell_shading(cell, LIGHT)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(value)
        set_run_font(r, size=9.2, bold=True, color=NAVY)

    for row_index, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "FAFBFC")
        for index, value in enumerate(values):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(value))
            set_run_font(r, size=9.2, color=INK)
    return table


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_props.append(color)
    run_props.append(underline)
    run.append(run_props)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def set_picture_alt(inline_shape, title: str, description: str):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("title", title)
    doc_pr.set("descr", description)


def add_figure(doc, image_path: Path, caption: str, alt: str, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shape = p.add_run().add_picture(str(image_path), width=Inches(width))
    set_picture_alt(shape, caption, alt)
    cp = doc.add_paragraph(style="Caption")
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.keep_with_next = False
    r = cp.add_run(caption)
    set_run_font(r, size=9, italic=True, color=MUTED)
    return shape


def add_page_field(paragraph):
    run = paragraph.add_run("Page ")
    set_run_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.append(begin)
    field_run.append(instr)
    field_run.append(separate)
    field_run.append(placeholder)
    field_run.append(end)
    paragraph._p.append(field_run)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    title._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)
    title_p_pr = title._element.get_or_add_pPr()
    title_border = title_p_pr.find(qn("w:pBdr"))
    if title_border is not None:
        title_p_pr.remove(title_border)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(15)
    subtitle.font.color.rgb = rgb(DARK_BLUE)
    subtitle._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    subtitle._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    subtitle.paragraph_format.space_after = Pt(12)

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)

    for style_name in ("Caliber Bullet", "Caliber Number"):
        if style_name not in styles:
            style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = rgb(INK)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_page(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_body_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    p = header.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run("CALIBER STUDIO")
    set_run_font(r, size=8.5, bold=True, color=MUTED)
    r = p.add_run("\tMASTER PLAN V4")
    set_run_font(r, size=8.5, color=MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_page_field(p)

    sec_pr = section._sectPr
    pg_num = sec_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sec_pr.append(pg_num)
    pg_num.set(qn("w:start"), "1")


def add_heading(doc, text: str, level: int = 1, new_page=False):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.page_break_before = new_page
    p.paragraph_format.keep_with_next = True
    return p


def add_contents_entry(doc, label: str, page: int | None = None, level: int = 0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22 * level)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    if page is not None:
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
    r = p.add_run(label)
    set_run_font(r, size=9.5, color=INK, bold=level == 0)
    if page is not None:
        r = p.add_run(f"\t{page}")
        set_run_font(r, size=9.5, color=MUTED)
    return p


def load_font(size: int, bold=False):
    path = FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(str(path), size)


def centered_text(draw, box, text, font, fill, max_width=None, line_gap=7):
    x0, y0, x1, y1 = box
    max_width = max_width or (x1 - x0 - 24)
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = word if not current else current + " " + word
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    heights = [draw.textbbox((0, 0), line, font=font)[3] for line in lines]
    total_h = sum(heights) + line_gap * max(0, len(lines) - 1)
    y = y0 + (y1 - y0 - total_h) / 2
    for line, height in zip(lines, heights):
        bbox = draw.textbbox((0, 0), line, font=font)
        width = bbox[2] - bbox[0]
        draw.text((x0 + (x1 - x0 - width) / 2, y), line, font=font, fill=fill)
        y += height + line_gap


def draw_box(draw, box, title, subtitle=None, fill="#FFFFFF", outline="#D6DCE4", accent=None):
    x0, y0, x1, y1 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    if accent:
        draw.rounded_rectangle((x0, y0, x0 + 12, y1), radius=6, fill=accent)
    if subtitle:
        centered_text(
            draw,
            (x0 + 20, y0 + 6, x1 - 14, y0 + (y1 - y0) * 0.54),
            title,
            load_font(25, True),
            "#0B1F33",
            line_gap=4,
        )
        centered_text(
            draw,
            (x0 + 25, y0 + (y1 - y0) * 0.58, x1 - 16, y1 - 6),
            subtitle,
            load_font(17),
            "#667085",
            line_gap=4,
        )
    else:
        centered_text(draw, (x0 + 20, y0 + 5, x1 - 16, y1 - 5), title, load_font(25, True), "#0B1F33")


def draw_arrow(draw, start, end, color="#2E74B5", width=5):
    draw.line((start, end), fill=color, width=width)
    x2, y2 = end
    x1, y1 = start
    dx, dy = x2 - x1, y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    tip = (x2, y2)
    left = (x2 - ux * 18 + px * 9, y2 - uy * 18 + py * 9)
    right = (x2 - ux * 18 - px * 9, y2 - uy * 18 - py * 9)
    draw.polygon((tip, left, right), fill=color)


def build_cover_banner(path: Path):
    width, height = 1500, 500
    image = Image.new("RGB", (width, height), "#" + NAVY)
    draw = ImageDraw.Draw(image)
    for y in range(height):
        ratio = y / height
        r = int(11 + 8 * ratio)
        g = int(31 + 24 * ratio)
        b = int(51 + 40 * ratio)
        draw.line((0, y, width, y), fill=(r, g, b))
    for x in range(0, width, 90):
        draw.line((x, 0, x, height), fill=(29, 63, 86), width=1)
    for y in range(0, height, 90):
        draw.line((0, y, width, y), fill=(29, 63, 86), width=1)

    nodes = [
        (160, 130, 38, TEAL),
        (400, 340, 28, BLUE),
        (660, 140, 46, GOLD),
        (900, 310, 34, TEAL),
        (1190, 125, 30, BLUE),
        (1350, 350, 44, TEAL),
    ]
    for first, second in zip(nodes, nodes[1:]):
        draw.line((first[0], first[1], second[0], second[1]), fill=(70, 132, 161), width=4)
    draw.line((nodes[0][0], nodes[0][1], nodes[3][0], nodes[3][1]), fill=(41, 91, 116), width=3)
    draw.line((nodes[2][0], nodes[2][1], nodes[5][0], nodes[5][1]), fill=(41, 91, 116), width=3)
    for x, y, radius, color in nodes:
        draw.ellipse((x - radius, y - radius, x + radius, y + radius), fill="#" + color, outline="#FFFFFF", width=5)
        draw.ellipse((x - 8, y - 8, x + 8, y + 8), fill="#FFFFFF")
    image.save(path)


def build_architecture_diagram(path: Path):
    image = Image.new("RGB", (1500, 920), "#F7F9FC")
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, True)
    draw.text((55, 36), "Caliber system architecture", font=title_font, fill="#" + NAVY)
    draw.text((55, 82), "One production control plane, multiple engines and asset providers", font=load_font(21), fill="#" + MUTED)

    draw_box(draw, (90, 150, 1410, 270), "Studio clients", "Web Studio | Godot plugin | Unity extension | Unreal plugin", fill="#FFFFFF", accent="#" + TEAL)
    draw_box(draw, (360, 350, 1140, 540), "Caliber Core", "Tasks | Events | Revisions | Changesets | Assets | Playtests | Builds | Permissions", fill="#E8EEF5", outline="#" + BLUE, accent="#" + BLUE)
    draw_box(draw, (80, 375, 300, 525), "OpenCode workers", "Director, code, scene, asset, test", fill="#FFFFFF", accent="#" + GOLD)
    draw_box(draw, (1200, 375, 1420, 525), "Engine adapters", "Typed native transactions", fill="#FFFFFF", accent="#" + TEAL)
    draw_box(draw, (90, 690, 430, 835), "Asset Foundry", "Tripo | Meshy | Blender | Libraries", fill="#FFFFFF", accent="#" + TEAL)
    draw_box(draw, (580, 690, 920, 835), "Evidence services", "Playtest | Performance | Validation | Publish", fill="#FFFFFF", accent="#" + BLUE)
    draw_box(draw, (1070, 690, 1410, 835), "Durable storage", "SQLite | Git | LFS | Artifact store", fill="#FFFFFF", accent="#" + GOLD)

    draw_arrow(draw, (750, 270), (750, 345))
    draw_arrow(draw, (300, 450), (350, 450), color="#" + GOLD)
    draw_arrow(draw, (1145, 450), (1190, 450), color="#" + TEAL)
    draw_arrow(draw, (610, 545), (300, 680), color="#" + TEAL)
    draw_arrow(draw, (750, 545), (750, 680), color="#" + BLUE)
    draw_arrow(draw, (890, 545), (1240, 680), color="#" + GOLD)
    image.save(path)


def build_asset_pipeline_diagram(path: Path):
    image = Image.new("RGB", (1500, 720), "#F7F9FC")
    draw = ImageDraw.Draw(image)
    draw.text((55, 35), "Asset Foundry production pipeline", font=load_font(34, True), fill="#" + NAVY)
    draw.text((55, 82), "Provider output is a candidate; Caliber makes it game-ready", font=load_font(21), fill="#" + MUTED)

    boxes = [
        ((70, 165, 330, 290), "1. Style-aware brief", "Purpose, references, scale, target budget", TEAL),
        ((420, 165, 680, 290), "2. Candidates", "Tripo, Meshy, upload, or library", BLUE),
        ((770, 165, 1030, 290), "3. Intake", "Download originals, metadata, cost, lineage", GOLD),
        ((1120, 165, 1380, 290), "4. Processing", "Scale, pivot, topology, UVs, textures", TEAL),
        ((1120, 470, 1380, 595), "5. Quality gates", "Visual, technical, engine, performance", BLUE),
        ((770, 470, 1030, 595), "6. Approval", "Approve, repair, replace, or reject", GOLD),
        ((420, 470, 680, 595), "7. Engine import", "Native resources, settings, stable identity", TEAL),
        ((70, 470, 330, 595), "8. Runtime evidence", "Usage, memory, LOD, draw calls, regressions", BLUE),
    ]
    for box, title, subtitle, color in boxes:
        draw_box(draw, box, title, subtitle, fill="#FFFFFF", accent="#" + color)
    for index in range(3):
        draw_arrow(draw, (boxes[index][0][2] + 8, 228), (boxes[index + 1][0][0] - 8, 228))
    draw_arrow(draw, (1250, 300), (1250, 460))
    for index in range(4, 7):
        draw_arrow(draw, (boxes[index][0][0] - 8, 532), (boxes[index + 1][0][2] + 8, 532))
    image.save(path)


def build_ui_diagram(path: Path):
    image = Image.new("RGB", (1500, 900), "#EEF2F6")
    draw = ImageDraw.Draw(image)
    draw.text((55, 34), "Caliber Web Studio", font=load_font(34, True), fill="#" + NAVY)
    draw.text((55, 80), "Prompt and agents on the left; directly editable game on the right", font=load_font(21), fill="#" + MUTED)

    frame = (55, 130, 1445, 840)
    draw.rounded_rectangle(frame, radius=22, fill="#FFFFFF", outline="#" + LINE, width=4)
    draw.rectangle((55, 130, 1445, 195), fill="#" + NAVY)
    draw.text((85, 150), "CALIBER", font=load_font(22, True), fill="#FFFFFF")
    draw.text((1180, 152), "Edit  |  Play  |  Review", font=load_font(18), fill="#DCE8F2")

    draw.rectangle((55, 195, 520, 840), fill="#F7F9FC", outline="#" + LINE, width=2)
    draw.text((85, 225), "CREATE  CODE  TASKS  AGENTS  CHANGES", font=load_font(16, True), fill="#" + BLUE)
    draw.rounded_rectangle((85, 275, 490, 430), radius=14, fill="#FFFFFF", outline="#" + LINE, width=2)
    draw.text((110, 298), "What should we build?", font=load_font(21, True), fill="#" + NAVY)
    draw.text((110, 345), "Create an atmospheric observatory", font=load_font(17), fill="#" + MUTED)
    draw.text((110, 375), "with traversal, one hazard, and a", font=load_font(17), fill="#" + MUTED)
    draw.text((110, 405), "stylized sci-fi material pass.", font=load_font(17), fill="#" + MUTED)
    draw.rounded_rectangle((85, 470, 490, 620), radius=14, fill="#FFFFFF", outline="#" + LINE, width=2)
    draw.text((110, 495), "Active work", font=load_font(20, True), fill="#" + NAVY)
    draw.text((110, 535), "Scene Agent     assembling zone 2", font=load_font(16), fill="#" + MUTED)
    draw.text((110, 570), "Asset Agent     3 candidates ready", font=load_font(16), fill="#" + MUTED)

    draw.rectangle((520, 195, 1445, 700), fill="#1B2B3A")
    draw.polygon(((620, 650), (950, 315), (1350, 650)), fill="#31495B")
    draw.polygon(((760, 650), (1020, 410), (1250, 650)), fill="#52758A")
    draw.ellipse((920, 360, 1035, 475), outline="#" + TEAL, width=6)
    draw.line((977, 330, 977, 505), fill="#" + TEAL, width=4)
    draw.line((890, 417, 1064, 417), fill="#" + TEAL, width=4)
    draw.text((555, 220), "LIVE EDITABLE VIEWPORT", font=load_font(18, True), fill="#FFFFFF")
    draw.rounded_rectangle((1160, 225, 1415, 480), radius=12, fill="#F7F9FC", outline="#" + LINE, width=2)
    draw.text((1185, 250), "INSPECTOR", font=load_font(17, True), fill="#" + BLUE)
    draw.text((1185, 292), "ObservatoryDoor", font=load_font(18, True), fill="#" + NAVY)
    draw.text((1185, 335), "Position   12, 0, -4", font=load_font(15), fill="#" + MUTED)
    draw.text((1185, 370), "Material   Brass_Aged", font=load_font(15), fill="#" + MUTED)
    draw.text((1185, 405), "Revision   18", font=load_font(15), fill="#" + MUTED)
    draw.text((1185, 440), "Owner      You", font=load_font(15), fill="#" + TEAL)

    draw.rectangle((520, 700, 1445, 840), fill="#F7F9FC", outline="#" + LINE, width=2)
    draw.text((550, 723), "ASSET FOUNDRY", font=load_font(17, True), fill="#" + BLUE)
    for i, label in enumerate(("Brief", "Candidates", "Processing", "Validation", "Library")):
        x = 550 + i * 170
        draw.rounded_rectangle((x, 765, x + 145, 815), radius=10, fill="#FFFFFF", outline="#" + LINE, width=2)
        centered_text(draw, (x, 765, x + 145, 815), label, load_font(15, True), "#" + NAVY)
    image.save(path)


def add_cover(doc: Document, banner: Path):
    section = doc.sections[0]
    configure_page(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PRODUCT + TECHNICAL MASTER PLAN")
    set_run_font(r, size=10.5, bold=True, color=TEAL)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    remove_paragraph_border(p)
    r = p.add_run("CALIBER STUDIO")
    set_run_font(r, size=30, bold=True, color=NAVY)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-native, multi-engine game creation")
    set_run_font(r, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run("Direct editing. Asynchronous agents. Production-ready assets. Immersive styles.")
    set_run_font(r, size=11, italic=True, color=MUTED)

    add_figure(
        doc,
        banner,
        "Caliber connects creative direction, agents, assets, and game engines.",
        "Abstract network of connected nodes representing human direction, agents, assets, and engine adapters.",
        width=6.5,
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(26)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Definitive Master Plan v4")
    set_run_font(r, size=12, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("July 21, 2026")
    set_run_font(r, size=10.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Web private alpha -> Godot -> Unity -> Unreal")
    set_run_font(r, size=10.5, bold=True, color=BLUE)


def build_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    cover_banner = ASSET_DIR / "cover_banner.png"
    architecture = ASSET_DIR / "architecture.png"
    asset_pipeline = ASSET_DIR / "asset_pipeline.png"
    ui_diagram = ASSET_DIR / "studio_ui.png"
    build_cover_banner(cover_banner)
    build_architecture_diagram(architecture)
    build_asset_pipeline_diagram(asset_pipeline)
    build_ui_diagram(ui_diagram)

    doc = Document()
    configure_styles(doc)
    bullet_num_id = add_numbering_definition(doc, "bullet")
    decimal_num_id = add_numbering_definition(doc, "decimal")
    evidence_num_id = add_numbering_definition(doc, "decimal")

    add_cover(doc, cover_banner)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_page(body_section)
    configure_body_header_footer(body_section)

    add_heading(doc, "Document control", 1)
    add_table(
        doc,
        ("Field", "Value"),
        (
            ("Document", "Caliber Studio Product and Technical Master Plan"),
            ("Version", "4.0"),
            ("Status", "Definitive working plan"),
            ("Date", "July 21, 2026"),
            ("Planning horizon", "Web private alpha through Godot, Unity, and Unreal adapters"),
            ("Primary outcome", "A polished, directly editable, agent-built Web3D vertical slice published to a URL"),
        ),
        (2700, 6660),
    )
    add_callout(
        doc,
        "Strategic decision",
        "Build Caliber as the AI production control plane around proven rendering and game engines. Prove the complete workflow on Web3D, then transfer it to Godot, Unity, and Unreal through native adapters.",
    )

    add_heading(doc, "Contents", 1, new_page=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("The headings below mirror Word's Navigation Pane for quick browsing.")
    set_run_font(r, size=9.5, color=MUTED, italic=True)
    for label, page in (
        ("1. Executive summary", 3),
        ("2. Product boundaries", 3),
        ("3. Target experience and first game", 4),
        ("4. Product interface", 5),
        ("5. Platform strategy", 6),
        ("6. System architecture", 7),
        ("7. Technology stack", 8),
        ("8. Web renderer strategy", 9),
        ("9. Style Packs and immersion", 9),
        ("10. Asset Foundry", 11),
        ("11. Multi-agent collaboration", 13),
        ("12. Direct editing and change safety", 13),
        ("13. Engine adapter contract", 14),
        ("14. Performance and smoothness", 15),
        ("15. Playtesting and validation", 16),
        ("16. Persistence, security, and source control", 16),
        ("17. Delivery roadmap", 18),
        ("18. Team and operating model", 18),
        ("19. Metrics", 19),
        ("20. Risks and decision gates", 20),
        ("21. Definitions of done", 21),
        ("Appendix A. Initial MCP tool surface", 22),
        ("Appendix B. Initial backlog", 22),
        ("Appendix C. Official technical sources", 23),
        ("Final product statement", 24),
    ):
        add_contents_entry(doc, label, page=page)

    add_heading(doc, "1. Executive summary", 1, new_page=True)
    add_body(
        doc,
        "Caliber Studio is an AI-native, multi-engine game production environment. It combines prompting, code, durable tasks, direct scene editing, asset generation, style systems, playtesting, performance control, and publishing in one continuous loop."
    )
    add_callout(
        doc,
        "Product promise",
        "A creator chooses a style, describes a game experience, watches specialized agents work asynchronously, directly edits the live scene at any time, validates quality and performance, and publishes or exports the result.",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    add_heading(doc, "The five product pillars", 2)
    for lead, text in (
        ("Create. ", "Prompting, plans, code, and agent collaboration remain visible and reviewable."),
        ("Edit. ", "The game viewport is directly editable with selection, gizmos, an inspector, undo, and stable object identities."),
        ("Generate. ", "Asset Foundry produces and processes art, models, textures, rigs, animation, and later audio."),
        ("Immerse. ", "Versioned Style Packs coordinate rendering, assets, camera, animation, atmosphere, feedback, and audio."),
        ("Ship. ", "Caliber runs playtests, performance checks, builds, release validation, and publishing."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)

    add_heading(doc, "Initial market promise", 2)
    add_body(
        doc,
        "The first credible promise is not a complete AAA-sized game from one prompt. It is an AAA-inspired level of polish within a controlled scope: a five-to-ten-minute Web3D vertical slice with a coherent world, strong assets, responsive interaction, stable frame pacing, and a complete gameplay loop."
    )

    add_heading(doc, "2. Product boundaries", 1)
    add_heading(doc, "What Caliber owns", 2)
    for item in (
        "Project understanding, game intent, and structured memory.",
        "Durable asynchronous task orchestration and agent permissions.",
        "Safe revisions, locks, changesets, approvals, and recovery.",
        "Asset generation, processing, lineage, licensing metadata, and quality gates.",
        "Style Packs, immersion profiles, playtests, performance budgets, and release evidence.",
        "A common engine-adapter contract and a consistent production interface.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "What engines own", 2)
    for item in (
        "Rendering, physics, audio playback, native resources, and platform packaging.",
        "Engine-specific scene and scripting formats.",
        "Native editor surfaces when Caliber connects to Godot, Unity, or Unreal.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Explicit non-goals", 2)
    for item in (
        "Building a new renderer or physics engine.",
        "Perfectly converting a finished game between unrelated engines.",
        "Treating every AI-generated asset as production-ready.",
        "Supporting four engines at production quality on day one.",
        "Promising photoreal hero art without technical-art and human review.",
        "Mixing Three.js and Babylon.js inside one shipped project.",
    ):
        add_bullet(doc, item, bullet_num_id)

    add_heading(doc, "3. Target experience and first game", 1)
    add_heading(doc, "Initial user", 2)
    add_body(
        doc,
        "The first user is a technically curious solo creator or small game team that wants a polished prototype without manually coordinating every code, content, environment, and testing task."
    )
    add_heading(doc, "Reference vertical slice", 2)
    add_table(
        doc,
        ("Dimension", "Target"),
        (
            ("Format", "Five-to-ten-minute third-person action and exploration game"),
            ("World", "One compact environment split into three streamed areas"),
            ("Gameplay", "One traversal or combat mechanic, one objective, one interaction system"),
            ("Opposition", "Two enemy or hazard types"),
            ("Presentation", "Stylized Atmospheric Adventure Style Pack, spatial audio, effects, polished camera"),
            ("Assets", "Modular environment kit, normal props, and a small set of hero assets"),
            ("Delivery", "Playable desktop-browser URL with WebGPU preference and WebGL 2 fallback"),
            ("Performance", "1080p, 60 FPS target on defined reference hardware"),
        ),
        (2200, 7160),
    )
    add_heading(doc, "Core creator journeys", 2)
    for lead, text in (
        ("Start from intent. ", "Define genre, camera, art direction, target hardware, scope, and the first complete gameplay loop."),
        ("Directly edit. ", "Click an object, inspect it, move it, replace its material or source asset, and undo safely."),
        ("Generate an asset. ", "Create a style-aware brief, compare candidates, process the winner, validate it, and import it."),
        ("Build an environment. ", "Decompose a world into a graybox, modular kit, hero landmarks, props, lighting, effects, audio, and navigation."),
        ("Collaborate with agents. ", "Run code, scene, asset, performance, and test work asynchronously with explicit ownership."),
        ("Playtest and repair. ", "Reproduce failures with captured input, screenshots, state, logs, and performance evidence."),
        ("Publish. ", "Build a reproducible release, validate it in a clean browser, and publish a preview URL."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)

    add_heading(doc, "4. Product interface", 1)
    add_figure(
        doc,
        ui_diagram,
        "Figure 1. The Caliber Web Studio operating model.",
        "Wireframe showing prompt, code, tasks, agents, and changes on the left; a large live editable viewport on the right; selected-object inspector overlay; and Asset Foundry tray below.",
    )
    add_heading(doc, "Primary surfaces", 2)
    for lead, text in (
        ("Left workspace. ", "Create, Code, Tasks, Agents, and Changes expose intent, execution, ownership, evidence, and review."),
        ("Right viewport. ", "Selection, transform gizmos, drag-to-place, Edit, Play, Review, performance overlays, and agent-change highlights."),
        ("Inspector. ", "Stable ID, revision, transform, components, mesh, materials, collision, scripts, semantic role, and active owner."),
        ("Asset Foundry tray. ", "Brief, Candidates, Compare, Processing, Validation, Library, and Usage."),
        ("Engine-native panel. ", "Godot, Unity, and Unreal retain their own viewport, hierarchy, inspector, and undo while Caliber adds tasks, agents, assets, and review."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)

    add_heading(doc, "5. Platform strategy", 1)
    add_table(
        doc,
        ("Order", "Platform", "Purpose", "Entry gate"),
        (
            ("1", "Web3D", "Prove the complete Caliber experience, direct editing, assets, agents, playtests, and URL publishing.", "Immediate"),
            ("2", "Godot", "Prove the workflow transfers to an open native editor and engine.", "Web alpha stable"),
            ("3", "Unity", "Serve established C# production workflows and mobile or desktop projects.", "Adapter pattern proven"),
            ("4", "Unreal", "Serve high-end C++ and Blueprint production with binary-asset safety.", "Demand and specialist capacity"),
            ("5", "Web2D", "Add a separate Phaser-style template using the same Core and Asset Foundry.", "Web3D foundation stable"),
        ),
        (700, 1500, 4560, 2600),
    )
    add_heading(doc, "Why Web first", 2)
    for item in (
        "Immediate preview without launching an external editor.",
        "Direct ownership of the editing experience and hot reload.",
        "Shareable builds through a URL.",
        "Straightforward browser automation, screenshots, input replay, and performance measurement.",
        "Native use of GLB and glTF assets.",
        "A controlled environment for proving human-agent concurrency.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_callout(
        doc,
        "Scope rule",
        "Web is the reference implementation, not a permanent ceiling. The production model must transfer to native engines without pretending that engine-native scenes are universally interchangeable.",
    )

    add_heading(doc, "6. System architecture", 1, new_page=True)
    add_figure(
        doc,
        architecture,
        "Figure 2. Caliber system architecture.",
        "Architecture diagram showing Studio clients above Caliber Core; OpenCode workers and engine adapters on either side; Asset Foundry, evidence services, and durable storage below.",
    )
    add_heading(doc, "Caliber Core responsibilities", 2)
    for lead, text in (
        ("Project service. ", "Projects, paths, engine and renderer capabilities, Style Pack versions, and target profiles."),
        ("Event service. ", "Immutable events with actor, source, time, task, and correlation identity."),
        ("Task service. ", "Dependency graphs, worker leases, retries, pause and resume, evidence, and changesets."),
        ("Change service. ", "Expected revisions, operations, diffs, validation, approval, application, and reversal."),
        ("Lock service. ", "Short user-edit leases, agent scopes, exclusive binary locks, and expiry."),
        ("Agent service. ", "OpenCode worker lifecycle, permissions, concurrency, tool scope, and cost."),
        ("Asset service. ", "Provider jobs, downloads, lineage, processing, validation, approval, and import."),
        ("Playtest and build services. ", "Launch, input, capture, reports, reproducible builds, and publishing."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)
    add_heading(doc, "Protocol split", 2)
    add_body(doc, "Caliber deliberately separates high-frequency editor communication from bounded agent commands and large binary artifacts.")
    add_table(
        doc,
        ("Channel", "Use"),
        (
            ("Authenticated WebSocket", "Selection, drag lifecycle, inspector updates, logs, agent status, and performance samples"),
            ("MCP", "Typed bounded project, scene, asset, playtest, validation, and build tools"),
            ("HTTP or file handles", "Models, textures, screenshots, videos, builds, and large logs"),
        ),
        (2800, 6560),
    )

    add_heading(doc, "7. Technology stack", 1)
    add_table(
        doc,
        ("Responsibility", "Technology", "Reason"),
        (
            ("Caliber Core", "Rust", "Reliable local service, concurrency, performance, and controlled memory use"),
            ("Web Studio", "TypeScript + React", "Product UI, browser APIs, and rapid iteration"),
            ("Web renderer", "Three.js or Babylon.js", "WebGPU/WebGL, glTF, PBR, custom styles, and direct scene control"),
            ("Godot adapter", "GDScript first", "Fast editor integration; native extension only after profiling"),
            ("Unity adapter", "C#", "Native Unity editor language"),
            ("Unreal adapter", "C++", "Native Unreal editor integration"),
            ("Agent workers", "OpenCode", "Coding and tool execution under Caliber control"),
            ("Local state", "SQLite", "Transactional local-first persistence"),
            ("Artifacts", "Content-addressed storage", "Immutable originals, deduplication, lineage, and cacheability"),
            ("Desktop shell", "Tauri later", "Optional packaging around the Web Studio and Rust Core"),
        ),
        (2300, 2500, 4560),
    )
    add_body(
        doc,
        "Electron is not required. The Studio can run in a browser with a local Rust bridge for safe filesystem and process access. Packaging the editor does not determine Web game performance; rendering, scene budgets, loading, and asset processing do."
    )

    add_heading(doc, "8. Web renderer strategy", 1, new_page=True)
    add_table(
        doc,
        ("Criterion", "Three.js", "Babylon.js"),
        (
            ("Best fit", "Highly customized, artistic, or experimental rendering", "Faster delivery of conventional game systems"),
            ("Rendering", "WebGPURenderer, WebGL 2 fallback, TSL and node materials", "WebGPU/WebGL, PBR, node materials, integrated engine systems"),
            ("Style flexibility", "Very high; Caliber shapes more of the stack", "High; verify freedom against engine conventions"),
            ("Game systems", "Caliber integrates more physics, navigation, animation, and tooling", "More integrated cameras, animation, particles, audio, and physics paths"),
            ("Implementation effort", "Higher surrounding-engine effort", "Lower for traditional mechanics"),
            ("Project rule", "One renderer per project", "One renderer per project"),
        ),
        (1900, 3730, 3730),
    )
    add_heading(doc, "Ten-day bakeoff", 2)
    add_body(
        doc,
        "The same attractive GLB environment, semantic Style Pack, direct-edit operations, compressed assets, camera, and performance scenario run in both renderers."
    )
    for item in (
        "Visual fidelity and Style Pack expressiveness.",
        "WebGPU and WebGL 2 fallback reliability.",
        "Selection, transforms, material editing, scene persistence, and undo integration.",
        "Loading, streaming, compressed assets, frame time, memory, and bundle size.",
        "Physics, animation, camera, audio, and gameplay integration effort.",
        "Playtest automation and team development speed.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_callout(
        doc,
        "Decision rule",
        "Ship one reference renderer in the first alpha. Add the second only after the complete product loop is stable and maintenance capacity is demonstrated.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    add_heading(doc, "9. Style Packs and immersion", 1)
    add_heading(doc, "Style Pack contract", 2)
    add_body(
        doc,
        "A Style Pack is a versioned production contract shared by humans, agents, Asset Foundry, renderers, native engines, and validation. It is not a collection of prompt adjectives."
    )
    for item in (
        "Visual references, permitted usage, shape language, palette, and contrast hierarchy.",
        "Surface and material language, edge treatment, damage rules, detail density, and texel density.",
        "Asset prompt templates, negative constraints, and approved material archetypes.",
        "Lighting, sky, fog, atmosphere, shaders, post-processing, and fallback rules.",
        "Camera, motion, animation timing, particles, audio, UI, and interaction feedback.",
        "Target profiles, performance limits, review criteria, and engine compatibility.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Renderer bindings", 2)
    add_body(
        doc,
        "The semantic pack remains renderer-neutral while executable details are adapter-specific. A Three.js binding can use TSL, node materials, and WebGPURenderer post-processing. A Babylon binding uses its PBR, node-material, and post-process systems. Godot, Unity, and Unreal later receive native bindings."
    )
    add_heading(doc, "Initial styles", 2)
    for lead, text in (
        ("Stylized Atmospheric Adventure. ", "The alpha pack: readable silhouettes, authored-looking surfaces, warm/cool lighting, expressive particles, smooth third-person camera, and environmental spatial audio."),
        ("Cinematic Science Fiction. ", "Metallic and composite surfaces, emissive accents, restrained bloom, dense atmosphere, and diegetic cues."),
        ("Painterly Dreamscape. ", "Simplified geometry, painterly textures, stylized depth, controlled palette, and experimental shader treatment."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)
    add_heading(doc, "Immersion profile", 2)
    for item in (
        "Camera perspective, movement, comfort, and transition rules.",
        "Input acceleration, smoothing, dead zones, responsiveness, and accessibility.",
        "Character animation transitions and interaction feedback.",
        "Environmental reactions, weather, atmosphere, spatial audio, and music transitions.",
        "Streaming and loading behavior, haptics where available, and diegetic versus screen-space UI.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_callout(
        doc,
        "Quality principle",
        "Immersion is a measured system. A visually impressive frame does not pass if input feels delayed, camera motion is uncomfortable, feedback is missing, audio is disconnected, or frame pacing is unstable.",
    )

    add_heading(doc, "10. Asset Foundry", 1)
    add_figure(
        doc,
        asset_pipeline,
        "Figure 3. Asset Foundry turns provider output into approved game assets.",
        "Eight-stage asset pipeline from style-aware brief through candidates, intake, processing, quality gates, approval, engine import, and runtime evidence.",
    )
    add_heading(doc, "Asset categories", 2)
    for item in (
        "Concept images, mood boards, sprites, UI, icons, decals, and skyboxes.",
        "Materials, tileable textures, static props, and modular environment kits.",
        "Characters, creatures, rigs, animation clips, and effects source assets.",
        "Sound effects, music, and voice as later provider-neutral categories.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Provider strategy", 2)
    for lead, text in (
        ("Tripo first. ", "Initial 3D provider for text, image, and multiview generation, texturing, low-poly processing, rigging, retargeting, and conversion."),
        ("Meshy benchmark. ", "First comparison and fallback provider across generation, retexturing, remeshing, rigging, animation, and conversion."),
        ("Blender processor. ", "Deterministic local conversion, cleanup, decimation, baking, validation, and turntable rendering."),
        ("Existing assets. ", "Uploads, studio libraries, purchased marketplace content, and open-license sources remain first-class."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)
    p = add_heading(doc, "Quality tiers", 2)
    p.paragraph_format.page_break_before = True
    add_table(
        doc,
        ("Tier", "Purpose", "Approval rule"),
        (
            ("Placeholder", "Fast graybox and mechanic validation", "Cannot ship"),
            ("Production", "Normal game use within technical and visual budgets", "Explicit approval after all required gates"),
            ("Hero", "Landmarks and focal characters or props", "Highest review bar; may require manual DCC work"),
        ),
        (1800, 4700, 2860),
    )
    add_heading(doc, "Quality gates", 2)
    for lead, text in (
        ("Geometry. ", "Parseability, orientation, scale, pivot, normals, tangents, topology, triangle budget, LODs, and collision."),
        ("UV and textures. ", "UV validity, texel density, seams, color space, PBR channels, dimensions, compression, and memory."),
        ("Materials. ", "Material count, plausible PBR values, supported shader features, transparency, and fallback."),
        ("Rig and animation. ", "Skeleton, bone budget, weights, deformation, root motion, clip naming, and retarget pose."),
        ("Visual. ", "Pinned Style Pack, silhouette, palette, detail density, gameplay-camera readability, and originality review."),
        ("Engine. ", "Import, save and reload, materials, animation, console errors, and runtime performance."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)
    add_heading(doc, "Lineage, licensing, and cost", 2)
    add_body(
        doc,
        "Every asset records the brief, references, prompts, provider, model version, seed when available, artifact hashes, generation cost, processing versions, human edits, source terms, approval identity, engine variants, and usage. Candidate generations remain in the artifact store; only approved sources and runtime assets enter source control."
    )

    add_heading(doc, "11. Multi-agent collaboration", 1, new_page=True)
    add_heading(doc, "Initial roles", 2)
    add_table(
        doc,
        ("Role", "Primary responsibility"),
        (
            ("Director", "Decompose intent into tasks, dependencies, scopes, and acceptance criteria"),
            ("Code Worker", "Implement gameplay or tools code in an isolated Git worktree"),
            ("Scene Worker", "Apply structured scene operations through an engine adapter"),
            ("Asset Worker", "Create briefs, request candidates, and run Asset Foundry processing"),
            ("Environment Worker", "Assemble modular kits and dress a controlled scene scope"),
            ("Test Worker", "Run builds and playtests; collect evidence without mutating production resources"),
            ("Performance Worker", "Measure budgets and propose optimizations"),
            ("Integrator", "Check revisions, conflicts, evidence, tests, and approved integration"),
            ("Repair Worker", "Implement the smallest fix for a reproduced failure"),
        ),
        (2200, 7160),
    )
    add_heading(doc, "Human editing has priority", 2)
    for step in (
        "The client requests a short lease when the user starts editing an object.",
        "Caliber rechecks the object revision before the transaction begins.",
        "Conflicting agent mutations wait while the user directly edits.",
        "The user commits; native undo history is preserved; the revision increments.",
        "Waiting agents receive the new state and either rebase or escalate.",
    ):
        add_numbered(doc, step, decimal_num_id)
    add_heading(doc, "Initial concurrency", 2)
    add_body(
        doc,
        "Start conservatively: one Director, one scene-writing worker, up to two disjoint code workers, one Asset Worker, one read-only Test or Performance Worker, and one Integrator. Expand only after collision, rebase, blocked-time, and integration metrics are healthy."
    )
    add_heading(doc, "Durable work record", 2)
    add_body(
        doc,
        "The source of truth is the task, resource scope, decisions, changeset, evidence, and approval. Chat remains useful context but is not the production record."
    )

    add_heading(doc, "12. Direct editing and change safety", 1)
    add_heading(doc, "Editor modes", 2)
    for lead, text in (
        ("Edit. ", "Selection and gizmos are active; scene changes persist; simulation is paused or controlled."),
        ("Play. ", "The game owns input; runtime state does not automatically become persistent editor state."),
        ("Review. ", "Proposed changes are overlaid, additions and removals are highlighted, and the user accepts or rejects."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)
    add_heading(doc, "Unified change model", 2)
    for item in (
        "Every resource has a stable Caliber ID, engine-native locator, revision, dependency edges, and active lease.",
        "Every mutation declares expected revisions and an idempotency key.",
        "Safe text and structured changes use optimistic concurrency.",
        "Binary assets and fragile shared scenes use exclusive ownership.",
        "Engine adapters group changes into native undo transactions or provide an explicit compensating reversal.",
        "Large binary data never moves through MCP.",
    ):
        add_bullet(doc, item, bullet_num_id)

    add_heading(doc, "13. Engine adapter contract", 1)
    add_heading(doc, "Required capabilities", 2)
    for item in (
        "Connect and authenticate; report project and adapter capabilities.",
        "Enumerate scenes and resources; report and inspect selection.",
        "Apply a transactional changeset with expected revisions and idempotency.",
        "Support native undo or a compensating reversal.",
        "Import an approved asset and retain stable Caliber identity.",
        "Launch and stop playtests; capture logs and screenshots.",
        "Validate and build the project.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Native-engine sequence", 2)
    for lead, text in (
        ("Godot. ", "GDScript editor plugin first; native viewport, selection, UndoRedo, GLB import, command-line validation, and desktop or Web export."),
        ("Unity. ", "C# editor package; selection, Undo, Asset Database, scene and prefab scopes, Play Mode, tests, and builds."),
        ("Unreal. ", "C++ editor plugin; native transactions, actor and asset inspection, Play In Editor, source control, One File Per Actor where useful, and binary locks."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)
    add_callout(
        doc,
        "No fake universal scene",
        "Caliber shares intent, stable identities, semantic roles, asset references, task history, and evidence. Detailed scenes remain renderer- or engine-native so advanced engine features are not reduced to a lowest common denominator.",
    )

    add_heading(doc, "14. Performance and smoothness", 1)
    add_heading(doc, "Initial target profile", 2)
    for item in (
        "Desktop browser on defined reference hardware.",
        "WebGPU preferred with WebGL 2 fallback.",
        "1080p reference viewport and 60 FPS target.",
        "Stable frame pacing, responsive input, and progressive loading.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Budget dimensions", 2)
    add_body(
        doc,
        "Each project profile calibrates budgets against actual reference hardware. Caliber does not hard-code invented universal limits."
    )
    for item in (
        "Visible triangles, objects, draw calls, materials, and shader variants.",
        "Texture memory, animation cost, physics bodies, particles, and audio voices.",
        "Initial download, time to first playable frame, streaming stalls, and JavaScript memory.",
        "Average and low-percentile frame time, CPU and GPU estimates, and input latency indicators.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Web optimization pipeline", 2)
    for item in (
        "GLB and glTF runtime interchange with supported mesh compression.",
        "KTX2 or Basis texture compression, mipmaps, and controlled texture sizes.",
        "LODs, instancing, culling, light baking, and shader variant control.",
        "Streamed zones, lazy assets, bundle splitting, compressed audio, and memory-aware unloading.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_callout(
        doc,
        "Release rule",
        "Smoothness is a release criterion, not a final optimization pass. Every integrated changeset may carry a performance delta, and visible quality reductions require review.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    add_heading(doc, "15. Playtesting and validation", 1)
    add_heading(doc, "Evidence loop", 2)
    for step in (
        "Build or launch an instrumented game target.",
        "Send recorded input and capture screenshots, state, logs, and performance.",
        "Evaluate deterministic assertions and visual checkpoints.",
        "Create a reproducible Repair task when the scenario fails.",
        "Implement the smallest scoped fix and repeat the scenario plus regressions.",
    ):
        add_numbered(doc, step, evidence_num_id)
    add_heading(doc, "Verification matrix", 2)
    add_table(
        doc,
        ("Area", "Required evidence"),
        (
            ("Core", "Unit tests, event replay, lease expiry, idempotency, permissions, and crash recovery"),
            ("Adapter", "Connect, selection, inspection, transaction, conflict, undo, reload, import, playtest, validation, and build"),
            ("Web Studio", "Browser interaction, gizmos, hot reload, sandbox security, and scene persistence"),
            ("Asset Foundry", "Provider contracts, retries, hashes, lineage, processor determinism, quality fixtures, license metadata, and cost limits"),
            ("Game", "Input scenario, state assertions, screenshots, console errors, frame-time thresholds, and clean-browser release load"),
        ),
        (2200, 7160),
    )

    add_heading(doc, "16. Persistence, security, and source control", 1)
    add_heading(doc, "Persistence", 2)
    for lead, text in (
        ("SQLite. ", "Projects, events, tasks, leases, revisions, locks, changesets, agent runs, assets, provider jobs, playtests, and builds."),
        ("Event log. ", "Append important state transitions before updating derived state for audit, replay, recovery, and UI streaming."),
        ("Artifact store. ", "Content-addressed immutable originals, processor caching, local storage first, and optional remote object storage."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)
    add_heading(doc, "Security and privacy", 2)
    for item in (
        "Core binds to loopback by default and clients use short-lived authentication.",
        "Provider keys remain in Core or the operating-system keychain.",
        "Game previews run in a sandbox without credentials or arbitrary host filesystem access.",
        "Imported files receive signature, type, size, path, and parser safety checks.",
        "Projects declare whether external providers, specific vendors, or team cloud storage are allowed.",
        "Agent shell and tool access remain restricted to approved project scope.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Source control", 2)
    for item in (
        "Code workers use isolated Git worktrees.",
        "Web and Godot favor isolated text-based resources.",
        "Unity coordinates scenes and prefabs carefully.",
        "Unreal uses source-control-aware binary locks and One File Per Actor where useful.",
        "Caliber never attempts a blind automatic merge of opaque binary assets.",
    ):
        add_bullet(doc, item, bullet_num_id)

    add_heading(doc, "17. Delivery roadmap", 1, new_page=True)
    add_table(
        doc,
        ("Phase", "Duration", "Primary deliverable", "Exit gate"),
        (
            ("0. Architecture and quality spike", "2 weeks", "Renderer bakeoff, Core event, direct transform, Tripo sample, Style Pack v0", "Reference renderer selected from evidence"),
            ("1. Web editor foundation", "4 weeks", "Projects, hierarchy, viewport, gizmos, inspector, undo, modes, style binding", "Small scene built without code"),
            ("2. Safe AI and multi-agent loop", "4 weeks", "OpenCode lifecycle, tasks, leases, revisions, changesets, MCP, recovery", "User and two disjoint agents without lost edits"),
            ("3. Asset Foundry alpha", "5 weeks", "Provider adapters, candidates, lineage, Blender, validation, approval, replacement", "Production asset replaces placeholder safely"),
            ("4. Immersion, playtest, performance, publish", "5 weeks", "Evidence, streaming, camera, audio, performance, release, preview URL", "New user completes reference flow"),
            ("5. Godot adapter", "6-8 weeks", "Native plugin, transactions, GLB import, playtest, validation, builds", "Shared Caliber concepts work natively"),
            ("6. Unity adapter", "8-10 weeks", "C# editor integration and production workflow", "Adapter contract passes"),
            ("7. Unreal adapter", "10-14 weeks", "C++ editor integration and binary-safe production workflow", "Adapter contract passes"),
        ),
        (2300, 1300, 3260, 2500),
    )
    add_heading(doc, "Schedule reality", 2)
    add_body(
        doc,
        "A focused team of three to four experienced engineers can target the Web private alpha in approximately 18-22 calendar weeks. A solo implementation is more realistically 30-44 focused weeks. Godot, Unity, and Unreal are not included in the Web alpha estimate."
    )
    add_heading(doc, "First 30 working days", 2)
    add_table(
        doc,
        ("Days", "Focus", "Observable success"),
        (
            ("1-5", "Initialize monorepo; draft protocol; build Three.js/Babylon scenes; target profile; Style Pack and immersion drafts", "Same GLB and style intent captured in both renderers"),
            ("6-10", "Choose renderer; pin manifest; implement style binding, selection, gizmos, inspector, undo, and persistence", "User edits survive refresh and undo"),
            ("11-15", "Rust Core, SQLite migrations, event replay, authentication, revisions, idempotent mutation", "Restart loses no history and duplicates no mutation"),
            ("16-20", "One OpenCode worker, Director, scene MCP tool, code task, worktree, changeset review", "Agent and user edit disjoint objects safely"),
            ("21-25", "Asset job schema, secret storage, Tripo job, output download, lineage, turntable, first QA report", "Candidate appears but cannot import before approval"),
            ("26-30", "Normalize GLB, pivot, scale, collision, simple budgets, import, placeholder replacement, performance capture", "Prompt-to-approved editable game asset works end to end"),
        ),
        (1100, 5260, 3000),
    )

    add_heading(doc, "18. Team and operating model", 1)
    add_heading(doc, "Minimum serious Web alpha team", 2)
    for item in (
        "Product and AI systems lead.",
        "Senior Web3D and graphics engineer.",
        "Rust and infrastructure engineer.",
        "Technical artist or tools artist, at least part time.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_body(
        doc,
        "Engine phases add specialists: a Godot tools engineer, a Unity editor-tools engineer, and an Unreal C++ editor-tools engineer. Technical art is an early requirement because API success and triangle counts cannot determine whether an asset is visually coherent, deformable, efficient, or ready to ship."
    )
    add_heading(doc, "Engineering operating principles", 2)
    for item in (
        "One monorepo initially with independently versioned protocol schemas.",
        "Engine adapters depend on contracts, not Core internals.",
        "One golden Web3D project exercises assets, character, animation, physics, UI, audio, loading, Style Pack, camera, and feedback.",
        "Small verified milestones, narrow tests, and explicit acceptance criteria.",
        "No engine expansion until the current adapter remains maintainable.",
    ):
        add_bullet(doc, item, bullet_num_id)

    add_heading(doc, "19. Metrics", 1)
    add_callout(
        doc,
        "North-star metric",
        "Weekly number of user-approved, playtested, performance-passing game improvements.",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    add_heading(doc, "Supporting measures", 2)
    for lead, text in (
        ("Creation. ", "Time to first playable scene, prompt-to-visible change, and asset brief-to-approved import."),
        ("Assets. ", "Approval rate, repair rate, variants per approval, cost per approved asset, import success, and style score."),
        ("Agents. ", "Completion, correction, conflict, rebase, blocked time, and duplicate-operation rates."),
        ("Reliability. ", "Crash-free sessions, recovery success, lost-edit count, build success, and adapter contract pass rate."),
        ("Game quality. ", "Frame-time pass rate, time to interactive, budget violations, runtime errors, playtest results, visual regressions, comfort failures, and audio-zone validation."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)

    add_heading(doc, "20. Risks and decision gates", 1, new_page=True)
    add_heading(doc, "Major risks", 2)
    add_table(
        doc,
        ("Risk", "Consequence", "Mitigation"),
        (
            ("Four engines too early", "No engine feels reliable", "Ship Web, then Godot, then expand"),
            ("Two Web renderers before alpha", "Double maintenance delays the product", "Bake off both; ship one; add the second later"),
            ("Universal scene abstraction", "Advanced engine features are lost", "Share semantic intent while keeping native scenes"),
            ("Inconsistent generated assets", "Game feels like random AI output", "Pinned Style Pack, references, variants, and technical-art review"),
            ("Provider output is technically weak", "Bad performance or deformation", "Processing and quality gates before import"),
            ("Web scene is attractive but slow", "Product misses its promise", "Continuous target profiles and scene budgets"),
            ("Agent overwrites the user", "Trust is lost", "User-priority leases, revisions, transactions, and undo"),
            ("Licensing is unclear", "Commercial risk", "Provenance, terms reference, review, and no silent redistribution"),
            ("Photorealism too early", "Unbounded art workload", "Start stylized and define explicit quality tiers"),
            ("Browser trust boundary is weak", "Credential or filesystem exposure", "Sandboxed preview, local Core, scoped tools, and keychain secrets"),
        ),
        (3000, 2760, 3600),
    )
    add_heading(doc, "Decision gates", 2)
    for lead, text in (
        ("Gate 1 - Renderer. ", "End of day 10: select Three.js or Babylon.js from visual, style, fallback, editing, loading, performance, bundle, systems, and productivity evidence."),
        ("Gate 2 - 3D provider. ", "Benchmark Tripo and Meshy on the same prop, modular piece, stylized character, and material brief before choosing defaults by asset class."),
        ("Gate 3 - Godot. ", "Start only after safe concurrency, provider-neutral assets, the adapter contract, playtest evidence, and stable Core are proven."),
        ("Gate 4 - Unity and Unreal. ", "Start only with user demand, maintenance capacity, a proven native-adapter pattern, and engine-specific expertise."),
    ):
        add_bullet(doc, text, bullet_num_id, lead=lead)

    add_heading(doc, "21. Definitions of done", 1, new_page=True)
    add_heading(doc, "Web technical alpha", 2)
    for item in (
        "Supported desktop browser, selected renderer behind the Web renderer contract, and directly editable viewport.",
        "Selection, transforms, materials, undo, save, and reload are stable.",
        "At least two agents work on disjoint tasks while human edits retain priority.",
        "Core survives restart; changesets are reviewable and reversible.",
        "One generated 3D asset passes the full pipeline and preserves lineage.",
        "One versioned Style Pack drives assets, renderer settings, camera, effects, and audio guidance.",
        "Reference visual, immersion, playtest, and performance checks pass.",
        "A reproducible release publishes to a preview URL.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Asset Foundry alpha", 2)
    for item in (
        "Provider jobs are asynchronous, resumable, cost-limited, and secret-safe.",
        "Originals are immutable and candidates are comparable.",
        "Geometry, materials, textures, collision, and engine import are validated.",
        "Production approval is explicit and asset replacement preserves scene references.",
        "Provider and processor failures produce actionable reasons.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "Godot adapter alpha", 2)
    for item in (
        "Native plugin connects to Core and uses native selection and undo.",
        "Revision conflicts are safe and approved GLB assets import.",
        "Playtests and validation use the shared evidence schema.",
        "A small project builds successfully for the selected target.",
    ):
        add_bullet(doc, item, bullet_num_id)

    add_heading(doc, "Appendix A. Initial MCP tool surface", 1, new_page=True)
    add_table(
        doc,
        ("Domain", "Initial tools"),
        (
            ("Project", "project.status; project.capabilities; project.search; resource.inspect; resource.dependencies"),
            ("Scene", "scene.list; scene.inspect; scene.selection; scene.apply_changeset; scene.validate; scene.save"),
            ("Asset", "asset.search; asset.inspect; asset.generate_candidates; asset.get_job; asset.process; asset.validate; asset.approve; asset.import; asset.replace_reference"),
            ("Playtest", "playtest.start; playtest.send_input; playtest.capture; playtest.read_state; playtest.stop; playtest.report"),
            ("Build", "build.development; build.release; build.report; publish.preview"),
        ),
        (1800, 7560),
    )
    add_body(
        doc,
        "All mutation tools require an approved task scope. High-risk operations can require human approval. Asset generation obeys project and task cost ceilings. Idempotency prevents duplicate provider jobs and duplicate scene changes."
    )

    add_heading(doc, "Appendix B. Initial backlog", 1)
    add_heading(doc, "P0 - Must prove", 2)
    for item in (
        "Web3D Studio and renderer decision.",
        "Versioned Style Pack and one production renderer binding.",
        "Direct selection, transform, material editing, undo, save, and reload.",
        "Rust Core persistence, revisions, user-priority leases, and changeset review.",
        "One OpenCode worker and one bounded scene tool.",
        "Tripo generation, GLB intake, validation, approval, and import.",
        "Camera, interaction, environmental audio, performance overlay, and local release build.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "P1 - Private alpha", 2)
    for item in (
        "Two disjoint workers, durable lineage, Blender processing, compression, and LODs.",
        "Playtest automation, preview hosting, Style Pack editor, visual reference captures, modular environment workflow, recovery UI, and cost controls.",
    ):
        add_bullet(doc, item, bullet_num_id)
    add_heading(doc, "P2 and later", 2)
    for item in (
        "Meshy production routing, second Web renderer, additional Style Packs, teams, cloud Core, Tauri Hub, Web2D, Godot, and audio generation.",
        "Unity, Unreal, photoreal profile, remote GPU workers, marketplace integrations, terrain pipelines, and console orchestration.",
    ):
        add_bullet(doc, item, bullet_num_id)

    add_heading(doc, "Appendix C. Official technical sources", 1, new_page=True)
    sources = (
        ("Three.js WebGPURenderer", "https://threejs.org/manual/en/webgpurenderer.html"),
        ("Three.js post-processing", "https://threejs.org/manual/en/post-processing.html"),
        ("Three.js WebXR", "https://threejs.org/manual/en/webxr-basics.html"),
        ("Three.js glTF loader", "https://threejs.org/docs/#examples/en/loaders/GLTFLoader"),
        ("Babylon.js WebGPU", "https://doc.babylonjs.com/setup/support/webGPU/"),
        ("Babylon.js glTF import", "https://doc.babylonjs.com/features/featuresDeepDive/importers/glTF/"),
        ("Babylon.js PBR", "https://doc.babylonjs.com/features/featuresDeepDive/materials/using/introToPBR/"),
        ("Khronos glTF", "https://www.khronos.org/gltf/"),
        ("Godot editor plugins", "https://docs.godotengine.org/en/stable/tutorials/plugins/editor/index.html"),
        ("Godot import plugins", "https://docs.godotengine.org/en/stable/tutorials/plugins/editor/import_plugins.html"),
        ("Godot 3D gizmos", "https://docs.godotengine.org/en/stable/tutorials/plugins/editor/3d_gizmos.html"),
        ("Godot Web export", "https://docs.godotengine.org/en/stable/tutorials/export/exporting_for_web.html"),
        ("Tripo OpenAPI", "https://docs.tripo3d.ai/get-started/introduction.html"),
        ("Tripo smart low poly", "https://docs.tripo3d.ai/mesh-editing/smart-low-poly-p-v2-0-20251225.html"),
        ("Tripo rigging", "https://docs.tripo3d.ai/animation/rig-v2-5-20260210.html"),
        ("Tripo terms", "https://www.tripo3d.ai/terms"),
        ("Meshy API", "https://docs.meshy.ai/en/api"),
        ("OpenCode MCP servers", "https://opencode.ai/docs/mcp-servers/"),
        ("OpenCode SDK", "https://opencode.ai/docs/sdk/"),
        ("OpenCode server", "https://opencode.ai/docs/server/"),
        ("MCP transports", "https://modelcontextprotocol.io/specification/2025-11-25/basic/transports"),
    )
    for label, url in sources:
        p = doc.add_paragraph(style="Caliber Bullet")
        apply_num(p, bullet_num_id)
        r = p.add_run(label + ": ")
        set_run_font(r, bold=True, color=INK)
        add_hyperlink(p, url, url)

    add_heading(doc, "Final product statement", 1)
    add_callout(
        doc,
        "Caliber Studio",
        "Caliber is the production system connecting human creative direction, direct visual editing, asynchronous agents, engine-native execution, high-quality assets, versioned styles, immersion, technical-art processing, performance, playtesting, review, and shipping.",
        fill=PALE_TEAL,
        accent=TEAL,
    )
    add_body(
        doc,
        "Web proves the complete loop first. Godot proves that it transfers to a native open engine. Unity and Unreal extend the same production system to larger professional workflows. The lasting moat is accumulated project understanding, asset lineage, Style Packs, immersion profiles, quality evidence, engine adapters, and a trustworthy human-agent production loop."
    )

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.core_properties.title = "Caliber Studio - Product and Technical Master Plan"
    doc.core_properties.subject = "AI-native multi-engine game creation platform"
    doc.core_properties.author = "Caliber Studio"
    doc.core_properties.keywords = "Caliber Studio, AI gaming, WebGPU, Three.js, Babylon.js, Godot, Unity, Unreal, multi-agent, Asset Foundry"
    doc.core_properties.comments = "Generated from the Caliber Studio definitive master plan v4."

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
