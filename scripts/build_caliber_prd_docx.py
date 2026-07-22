from __future__ import annotations

from pathlib import Path
from typing import Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import build_caliber_master_plan_docx as base


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / ".docx_build" / "prd_assets"
OUTPUT = ROOT / "CALIBER_STUDIO_PRD.docx"

VERSION = "1.1"
DATE = "July 21, 2026"

# Populated after the first render so the static contents page remains reliable
# in Word, LibreOffice, and preview renderers that do not update TOC fields.
CONTENTS_PAGES: dict[str, int] = {
    "1. Executive summary": 3,
    "2. Problem and opportunity": 4,
    "3. Product strategy and principles": 5,
    "4. Users and jobs to be done": 6,
    "5. Target experience and information architecture": 7,
    "6. Release scope and prioritization": 8,
    "7. Functional requirements": 9,
    "8. Critical user journeys": 14,
    "9. Product rules and domain model": 17,
    "10. Non-functional requirements": 18,
    "11. Trust, security, and privacy": 19,
    "12. Analytics and success metrics": 19,
    "13. Delivery and rollout plan": 21,
    "14. Release acceptance criteria": 22,
    "15. Dependencies, risks, and mitigations": 23,
    "16. Open decisions and decision gates": 24,
    "Appendix A. Requirement index by release": 25,
    "Appendix B. Product glossary": 25,
}


def configure_prd_header_footer(section):
    base.configure_page(section)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False

    header = section.header
    p = header.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run("CALIBER STUDIO")
    base.set_run_font(run, size=8.5, bold=True, color=base.MUTED)
    run = p.add_run("\tPRODUCT REQUIREMENTS DOCUMENT V1.1")
    base.set_run_font(run, size=8.5, color=base.MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    base.add_page_field(p)

    sec_pr = section._sectPr
    pg_num = sec_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sec_pr.append(pg_num)
    pg_num.set(qn("w:start"), "1")


def configure_prd_styles(doc: Document):
    # Resolves the standard_business_brief preset exactly, with named title and
    # requirement-table overrides reused throughout this PRD.
    base.configure_styles(doc)
    title = doc.styles["Title"]
    title.font.size = Pt(27)
    title.font.color.rgb = base.rgb(base.NAVY)
    title.paragraph_format.space_after = Pt(5)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = base.rgb(base.DARK_BLUE)
    subtitle.paragraph_format.space_after = Pt(14)


def add_metadata_line(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(f"{label}: ")
    base.set_run_font(run, size=10, bold=True, color=base.NAVY)
    run = p.add_run(value)
    base.set_run_font(run, size=10, color=base.INK)
    return p


def add_memo_masthead(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("PRODUCT REQUIREMENTS DOCUMENT")
    base.set_run_font(run, size=10.5, bold=True, color=base.TEAL)

    p = doc.add_paragraph(style="Title")
    base.remove_paragraph_border(p)
    run = p.add_run("CALIBER STUDIO")
    base.set_run_font(run, size=27, bold=True, color=base.NAVY)

    p = doc.add_paragraph(style="Subtitle")
    run = p.add_run("Web-first AI game production workspace")
    base.set_run_font(run, size=14, color=base.DARK_BLUE)

    add_metadata_line(doc, "Version", VERSION)
    add_metadata_line(doc, "Status", "Founder-ready product specification for private alpha")
    add_metadata_line(doc, "Date", DATE)
    add_metadata_line(doc, "Owner", "Founder / Product")
    add_metadata_line(doc, "Audience", "Product, design, engineering, technical art, and engine integration")
    add_metadata_line(doc, "Initial release", "Web3D private alpha; Godot is the next native-engine adapter")

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)

    base.add_callout(
        doc,
        "Product decision",
        "Caliber is not a replacement rendering engine. It is the AI-native production studio that connects human creative direction, direct visual editing, asynchronous agents, high-quality asset workflows, playtesting, performance evidence, and engine-native execution.",
        fill=base.PALE_TEAL,
        accent=base.TEAL,
    )

    base.add_heading(doc, "Product definition", 2)
    base.add_body(
        doc,
        "Caliber Studio lets a creator describe a game experience on the left, watch specialized agents work, and directly edit the playable world on the right. The first release proves the full loop on Web3D; later adapters bring the same production model into Godot, Unity, and Unreal without flattening their native capabilities."
    )
    base.add_heading(doc, "Private-alpha outcome", 2)
    base.add_body(
        doc,
        "A target user can create a coherent five-to-ten-minute Web3D vertical slice, edit scene objects while agents work on disjoint tasks, approve a generated and processed 3D asset, pass playtest and performance gates, and publish a reproducible preview URL."
    )


def add_requirement_table(
    doc: Document,
    rows: Sequence[tuple[str, str, str, str]],
):
    widths = (900, 4260, 800, 3400)
    table = doc.add_table(rows=1, cols=4)
    base.set_table_geometry(table, widths)
    header = table.rows[0]
    base.mark_header_row(header)
    base.prevent_row_split(header)
    labels = ("ID", "Requirement", "Priority", "Acceptance signal")
    for index, label in enumerate(labels):
        cell = header.cells[index]
        base.set_cell_shading(cell, base.LIGHT)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(label)
        base.set_run_font(run, size=8.7, bold=True, color=base.NAVY)

    priority_colors = {"P0": base.RED, "P1": base.GOLD, "P2": base.BLUE}
    for row_index, values in enumerate(rows):
        row = table.add_row()
        base.prevent_row_split(row)
        if row_index % 2 == 1:
            for cell in row.cells:
                base.set_cell_shading(cell, "FAFBFC")
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = 1.0
            color = priority_colors.get(value, base.INK) if index == 2 else base.INK
            run = p.add_run(value)
            base.set_run_font(
                run,
                size=8.65,
                color=color,
                bold=index in (0, 2),
            )
    return table


def add_requirement_section(
    doc: Document,
    title: str,
    lead: str,
    rows: Sequence[tuple[str, str, str, str]],
):
    base.add_heading(doc, title, 2)
    p = base.add_body(doc, lead)
    p.paragraph_format.keep_with_next = True
    return add_requirement_table(doc, rows)


def add_contents(doc: Document):
    base.add_heading(doc, "Contents", 1, new_page=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("The headings use real Word styles and mirror the Navigation Pane.")
    base.set_run_font(run, size=9.5, italic=True, color=base.MUTED)

    entries = (
        "1. Executive summary",
        "2. Problem and opportunity",
        "3. Product strategy and principles",
        "4. Users and jobs to be done",
        "5. Target experience and information architecture",
        "6. Release scope and prioritization",
        "7. Functional requirements",
        "8. Critical user journeys",
        "9. Product rules and domain model",
        "10. Non-functional requirements",
        "11. Trust, security, and privacy",
        "12. Analytics and success metrics",
        "13. Delivery and rollout plan",
        "14. Release acceptance criteria",
        "15. Dependencies, risks, and mitigations",
        "16. Open decisions and decision gates",
        "Appendix A. Requirement index by release",
        "Appendix B. Product glossary",
    )
    for label in entries:
        base.add_contents_entry(doc, label, page=CONTENTS_PAGES.get(label))


def add_front_matter(doc: Document):
    add_memo_masthead(doc)
    add_contents(doc)


def build_document():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    studio_ui = ASSET_DIR / "studio_ui.png"
    asset_pipeline = ASSET_DIR / "asset_pipeline.png"
    base.build_ui_diagram(studio_ui)
    base.build_asset_pipeline_diagram(asset_pipeline)

    doc = Document()
    configure_prd_styles(doc)
    configure_prd_header_footer(doc.sections[0])

    bullet_num_id = base.add_numbering_definition(doc, "bullet")
    decimal_num_id = base.add_numbering_definition(doc, "decimal")

    add_front_matter(doc)

    base.add_heading(doc, "1. Executive summary", 1, new_page=True)
    base.add_callout(
        doc,
        "Product promise",
        "Describe, delegate, directly edit, play, improve, and ship without losing creative control or leaving the production context.",
        fill=base.PALE_TEAL,
        accent=base.TEAL,
    )
    base.add_body(
        doc,
        "Caliber Studio is an AI-native game production workspace. It combines the speed of coding agents with the immediacy of a live game editor: prompting, plans, code, tasks, and diffs remain visible on the left while a persistent playable viewport, scene hierarchy, selection tools, and inspector remain available on the right."
    )
    base.add_body(
        doc,
        "Caliber uses existing game engines and renderers instead of attempting to replace them. Its product value is the production control plane around those engines: project understanding, agent orchestration, safe changes, asset generation and processing, style coherence, playtesting, performance budgets, approvals, evidence, and publishing."
    )
    base.add_heading(doc, "Initial product", 2)
    for lead, text in (
        ("Platform. ", "A browser-based Web3D Studio connected to a local Caliber Core; no Electron dependency is required for the first proof."),
        ("Game target. ", "A polished, stylized five-to-ten-minute 3D vertical slice for desktop browsers."),
        ("Renderer. ", "Three.js and Babylon.js complete a short evidence-based bakeoff; one becomes the production renderer for alpha."),
        ("Agent runtime. ", "OpenCode is the first coding-worker backend behind a provider-neutral internal interface."),
        ("Assets. ", "A provider-neutral Asset Foundry begins with a Tripo integration and compares Meshy before provider defaults are expanded."),
        ("Native engines. ", "Godot follows after the Web loop is stable; Unity and Unreal follow only when demand and adapter capacity justify them."),
    ):
        base.add_bullet(doc, text, bullet_num_id, lead=lead)
    base.add_heading(doc, "What success means", 2)
    base.add_body(
        doc,
        "The private alpha succeeds when users repeatedly create approved, playable, performance-passing improvements with less context switching and no silent loss of human edits. A one-prompt full AAA game is not an alpha claim; the product instead proves AAA-inspired polish inside a controlled vertical-slice scope."
    )

    base.add_heading(doc, "2. Problem and opportunity", 1)
    base.add_heading(doc, "User problem", 2)
    base.add_body(
        doc,
        "General-purpose coding agents can change files, but they do not provide a trustworthy game-production loop. Game editors expose scenes and objects, but they are not designed around durable asynchronous agent work. Asset generators produce files, but those files often arrive without consistent style, import readiness, provenance, performance validation, or safe replacement."
    )
    base.add_heading(doc, "Fragmented workflow", 2)
    for item in (
        "Creative intent is split across chat, code, scene files, art references, and human memory.",
        "The creator cannot safely manipulate the live world while multiple agents change nearby code or content.",
        "Generated assets require manual downloading, repair, scaling, pivot correction, materials, collision, optimization, and import.",
        "Visual quality, interaction quality, performance, and playtest evidence are evaluated late or inconsistently.",
        "Moving from a Web prototype to a native engine usually loses task history, asset lineage, and structured intent.",
    ):
        base.add_bullet(doc, item, bullet_num_id)
    base.add_heading(doc, "Product opportunity", 2)
    base.add_body(
        doc,
        "Caliber can become the place where game intent remains durable while code, scenes, and assets evolve. Its defensible advantage is not a chat box. It is the accumulated project model, safe human-agent collaboration, direct manipulation, reusable Style Packs, asset lineage, playtest evidence, engine adapters, and quality data produced by every iteration."
    )
    base.add_heading(doc, "Why Web first", 2)
    for item in (
        "The editor and game can share an immediate visual feedback loop with fast reloads and simple preview distribution.",
        "A URL is the lowest-friction way to test, share, and collect feedback on a vertical slice.",
        "Web3D is powerful enough to prove selection, scene editing, assets, agents, immersion, performance budgets, and publishing.",
        "A Web-first product avoids committing the studio shell to Electron before a native wrapper is justified.",
        "The shared Caliber concepts can later transfer to Godot without pretending that Web, Godot, Unity, and Unreal scenes are identical.",
    ):
        base.add_bullet(doc, item, bullet_num_id)

    base.add_heading(doc, "3. Product strategy and principles", 1)
    base.add_heading(doc, "Product thesis", 2)
    base.add_body(
        doc,
        "Caliber wins by shortening the loop from creative intent to verified playable result while increasing user control. The product must make automation feel observable and reversible, not magical and fragile."
    )
    base.add_heading(doc, "Product principles", 2)
    principles = (
        ("The game stays visible. ", "The right-side viewport persists while the creator prompts, reviews code, or inspects agent work."),
        ("Direct manipulation is first-class. ", "Clicking, moving, rotating, scaling, and editing an object must not require a prompt."),
        ("Human edits have priority. ", "Agents cannot silently overwrite a resource the user is actively editing."),
        ("Every important change is reviewable. ", "Mutations produce a scoped changeset, evidence, status, and reversal path."),
        ("Agents work from bounded tasks. ", "Chat may create context, but durable tasks, scopes, decisions, and evidence are the production record."),
        ("Assets are a pipeline, not a download. ", "Generation is followed by processing, validation, approval, import, lineage, and replacement."),
        ("Style is versioned product data. ", "A Style Pack coordinates assets, materials, rendering, camera, effects, feedback, and audio direction."),
        ("Smoothness is a release requirement. ", "Performance budgets and frame evidence are enforced throughout production."),
        ("Engines remain native. ", "Caliber shares intent and identity while engine adapters preserve native scenes, resources, undo, and build behavior."),
        ("Start narrow, prove, then expand. ", "Web first, Godot next, then Unity and Unreal when the adapter contract and product demand are proven."),
    )
    for lead, text in principles:
        base.add_bullet(doc, text, bullet_num_id, lead=lead)
    base.add_heading(doc, "Explicit non-goals for private alpha", 2)
    for item in (
        "Building a new general-purpose renderer, physics engine, audio engine, or AAA engine.",
        "Generating a complete production-scale AAA game from one prompt.",
        "Supporting Web, Godot, Unity, and Unreal at production quality on day one.",
        "Lossless conversion of a finished game between unrelated engines.",
        "Real-time multi-user editing, cloud organizations, marketplace distribution, or console certification.",
        "Accepting arbitrary generated assets as production-ready without technical-art and human approval.",
        "Pure no-code positioning; the initial user can inspect plans, diffs, code, and technical quality evidence.",
    ):
        base.add_bullet(doc, item, bullet_num_id)

    base.add_heading(doc, "4. Users and jobs to be done", 1)
    base.add_heading(doc, "Primary persona: technical creative", 2)
    base.add_body(
        doc,
        "A solo creator, technical designer, or indie developer who can reason about game systems but wants dramatically faster iteration. This user values direct control, can inspect code or diffs when necessary, and cares more about a playable, coherent result than the novelty of generation."
    )
    base.add_heading(doc, "Secondary personas", 2)
    persona_rows = (
        ("Small-studio lead", "Coordinates several disciplines and needs safe asynchronous progress, review, and a visible playable state."),
        ("Technical artist", "Turns generated or sourced content into efficient, coherent, engine-ready production assets."),
        ("Gameplay engineer", "Builds mechanics and integrations while relying on scoped tasks, tests, diffs, and native source control."),
        ("Designer / world builder", "Manipulates scenes directly, tunes values, tests feel, and delegates structured changes without living in code."),
        ("Engine tools engineer", "Maintains the Godot, Unity, or Unreal adapter and validates native transactions, imports, playtests, and builds."),
    )
    base.add_table(doc, ("Persona", "Primary need"), persona_rows, (2200, 7160))
    base.add_heading(doc, "Jobs to be done", 2)
    jobs = (
        ("When I have a game concept,", "help me reach a playable, visually coherent scene quickly so I can judge the idea instead of assembling infrastructure."),
        ("When I see something wrong in the world,", "let me click it and change it immediately so prompting is optional, not mandatory."),
        ("When several tasks can run independently,", "let specialized agents work asynchronously while I keep editing and playing."),
        ("When I need a new asset,", "create or source candidates, make them engine-ready, and preserve provenance so I can approve with confidence."),
        ("When the game feels inconsistent,", "apply a shared style and immersion system so assets, lighting, camera, effects, and sound reinforce one direction."),
        ("When a change fails,", "capture enough evidence to repair it without reconstructing the entire context."),
        ("When I am ready to share,", "prove that the game builds, loads, plays, and meets its performance target before publishing a preview."),
    )
    base.add_table(doc, ("Situation", "Desired outcome"), jobs, (2500, 6860))

    base.add_heading(doc, "5. Target experience and information architecture", 1, new_page=True)
    base.add_heading(doc, "Core experience", 2)
    base.add_body(
        doc,
        "The main studio is a resizable two-sided workspace. The left side supports intention, delegation, code, tasks, assets, and review. The right side supports direct interaction with the current game. Neither side is merely a preview of the other; together they form the production loop."
    )
    base.add_figure(
        doc,
        studio_ui,
        "Caliber Studio keeps prompting and production activity beside a persistent, directly editable game world.",
        "Wireframe of Caliber Studio with prompt, plan, task, and code areas on the left and an editable game viewport with hierarchy and inspector on the right.",
        width=6.35,
    )
    base.add_heading(doc, "Primary navigation", 2)
    nav_rows = (
        ("Create", "Prompt, plan, code, task queue, agent activity, diffs, approvals, and recovery."),
        ("Scene", "Hierarchy, viewport, selection, gizmos, inspector, modes, environment, and scene history."),
        ("Content", "Asset briefs, candidates, processing, validation, approvals, library, replacement, and Style Packs."),
        ("Verify", "Playtests, captured evidence, runtime errors, performance budgets, regressions, and repair tasks."),
        ("Ship", "Development build, release validation, build report, preview publishing, and export history."),
    )
    base.add_table(doc, ("Area", "Responsibilities"), nav_rows, (1700, 7660))
    base.add_heading(doc, "Editor modes", 2)
    for lead, text in (
        ("Edit. ", "Selection and persistent scene mutations are active. Simulation is paused or explicitly controlled."),
        ("Play. ", "The game owns input. Runtime changes do not silently become editor state."),
        ("Review. ", "Proposed additions, removals, and modifications are overlaid for accept, reject, or partial approval."),
    ):
        base.add_bullet(doc, text, bullet_num_id, lead=lead)
    base.add_heading(doc, "Status language", 2)
    base.add_body(
        doc,
        "The default experience uses product language such as Planning, Building mechanic, Generating candidates, Processing asset, Waiting for approval, Testing, Repair needed, Ready to play, and Ready to publish. Raw model and tool events remain available in a details view."
    )

    base.add_heading(doc, "6. Release scope and prioritization", 1)
    base.add_heading(doc, "Priority definitions", 2)
    priority_rows = (
        ("P0", "Must prove", "Required for the Web technical alpha and the core product thesis."),
        ("P1", "Private alpha", "Required before inviting a small external user cohort."),
        ("P2", "Post-alpha", "Valuable expansion after stability, usage evidence, and team capacity exist."),
    )
    base.add_table(doc, ("Priority", "Meaning", "Release rule"), priority_rows, (1100, 2200, 6060))
    base.add_heading(doc, "P0: Web technical alpha", 2)
    for item in (
        "Browser-based Web3D Studio connected to a durable local Core.",
        "Persistent right-side viewport with selection, transform gizmos, inspector edits, undo, save, and reload.",
        "Prompt, plan, code, changes, task, and review surfaces on the left.",
        "One coding worker, one scene-writing worker, one Asset Worker, and read-only test or performance work.",
        "Stable resource IDs, revisions, leases, idempotent changesets, user-priority conflict rules, and recovery.",
        "One versioned Style Pack and one renderer binding selected by the Three.js/Babylon.js bakeoff.",
        "One provider integration, candidate comparison, GLB intake, processing, validation, approval, and import.",
        "Deterministic playtest recording, synchronized frame traces, animation diagnosis, repair replay, performance evidence, release build, and local preview publishing.",
    ):
        base.add_bullet(doc, item, bullet_num_id)
    base.add_heading(doc, "P1: Private alpha", 2)
    for item in (
        "Two simultaneous disjoint writing agents with collision and recovery metrics.",
        "Blender-backed processing, compression, LOD generation, collision profiles, and stronger technical-art validation.",
        "Preview hosting, shareable feedback, recovery UI, cost controls, style editor, reference captures, and visual regression checks.",
        "A polished onboarding project that proves mechanics, animation, environment, audio, performance, and publishing end to end.",
    ):
        base.add_bullet(doc, item, bullet_num_id)
    base.add_heading(doc, "P2: Native engine expansion", 2)
    for item in (
        "Godot editor plugin and adapter contract implementation.",
        "Optional Tauri desktop hub when local distribution, keychain access, deep links, or process management justify it.",
        "Unity and Unreal adapters after demand, adapter stability, binary-safe workflows, and engine expertise are proven.",
        "Cloud teams, additional providers and Style Packs, marketplace integrations, remote workers, and console workflows.",
    ):
        base.add_bullet(doc, item, bullet_num_id)

    base.add_heading(doc, "7. Functional requirements", 1)
    base.add_body(
        doc,
        "All P0 requirements are release-blocking for the Web technical alpha unless a documented decision gate explicitly changes scope. P1 requirements are required for the external private-alpha cohort. Acceptance signals describe observable product behavior, not implementation alone."
    )

    add_requirement_section(
        doc,
        "7.1 Projects and studio shell",
        "The studio must open quickly into a durable project state and keep production context visible without trapping the project inside Caliber.",
        (
            ("PRJ-001", "Create a new Caliber Web3D project from the approved starter and open an existing Caliber project.", "P0", "A valid project launches into the studio and reaches a playable state without manual setup."),
            ("PRJ-002", "Persist project metadata, current scene, tasks, revisions, approvals, asset jobs, and build history across restarts.", "P0", "Restart restores the last accepted production state with no lost accepted edit."),
            ("PRJ-003", "Show project health, renderer, target profile, Core connection, adapter capabilities, and current build state.", "P0", "The user can identify a disconnected or invalid subsystem without opening developer tools."),
            ("PRJ-004", "Keep project source and content in normal readable folders with versionable structured files where practical.", "P0", "The project can be inspected and backed up outside Caliber; Core metadata has an export path."),
            ("PRJ-005", "Record accepted changes as revisions and expose history, evidence, author, task, and reversal status.", "P0", "A user can identify what changed and restore the prior accepted state."),
            ("UI-001", "Provide a resizable left production workspace and persistent right game workspace.", "P0", "Changing left-side modes does not unmount or reset the live game viewport."),
            ("UI-002", "Expose Prompt, Tasks, Code, Changes, Content, Verify, and Ship without overwhelming first-run users.", "P1", "Primary onboarding journey completes without requiring hidden developer-only navigation."),
            ("UI-003", "Surface concise human-readable status with optional raw agent and tool details.", "P0", "Every running task has a comprehensible current state and last meaningful event."),
            ("UI-004", "Support keyboard navigation, visible focus, scalable text, and inspector alternatives for pointer-only scene operations.", "P1", "Core workflows pass the product accessibility checklist at supported zoom levels."),
            ("UI-005", "Inspect and edit supported project code and text files in the left workspace with save, diagnostics, diff, and revision integration.", "P0", "A manual code edit saves to normal project source, appears in Changes, triggers the relevant rebuild or reload, and survives restart."),
        ),
    )

    add_requirement_section(
        doc,
        "7.2 Direct scene editing",
        "The viewport is an editor, not a screenshot. Direct user input must remain safe while agents work asynchronously.",
        (
            ("EDT-001", "Select an object from the viewport or hierarchy and keep both selections synchronized.", "P0", "Clicking a visible object highlights it and opens the correct inspector within 100 ms on the reference scene."),
            ("EDT-002", "Assign every editable resource a stable Caliber ID plus an engine-native locator and revision.", "P0", "Save, reload, asset replacement, and agent changes preserve the intended identity."),
            ("EDT-003", "Move, rotate, and scale selected objects with gizmos and numeric inspector fields.", "P0", "A transform survives save, reload, play-mode transitions, and Core restart."),
            ("EDT-004", "Edit supported component, material, light, camera, physics, and gameplay properties in a typed inspector.", "P0", "Invalid values are rejected with a useful message; accepted values update the viewport."),
            ("EDT-005", "Group a direct edit into a transaction with undo, redo, revision, and changeset history.", "P0", "Undo restores the exact previous accepted state without corrupting adjacent changes."),
            ("EDT-006", "Separate Edit, Play, and Review modes and make the active mode unmistakable.", "P0", "Runtime input never silently persists a scene mutation outside an explicit capture flow."),
            ("EDT-007", "Give active human edits priority over agent leases and require agents to rebase or request review on conflict.", "P0", "A controlled collision test produces no silent overwrite and an actionable conflict state."),
            ("EDT-008", "Allow approved asset replacement without breaking scene references or object-level overrides.", "P1", "Replacing a placeholder keeps transform, semantic role, and supported override data."),
        ),
    )

    add_requirement_section(
        doc,
        "7.3 AI tasks and asynchronous collaboration",
        "Agents operate through durable tasks, bounded tools, resource scopes, and reviewable changesets. MCP is a tool boundary, not the game engine.",
        (
            ("AGT-001", "Convert a user request into a visible plan or scoped task before production mutation when the change is non-trivial.", "P0", "The user can see objective, scope, owner, dependencies, and expected evidence."),
            ("AGT-002", "Run specialized coding, scene, asset, test, performance, and integration roles asynchronously.", "P0", "At least one writer and one read-only verifier can progress without blocking the user."),
            ("AGT-003", "Grant each task explicit project, resource, tool, network, cost, and time scope.", "P0", "An out-of-scope mutation is denied and recorded rather than executed."),
            ("AGT-004", "Lease mutable resources with expiry, renewal, expected revision, and user-priority semantics.", "P0", "Expired or conflicting leases cannot create a silent last-write-wins overwrite."),
            ("AGT-005", "Support start, cancel, retry, supersede, block, and resume for durable tasks.", "P0", "A Core restart or worker crash does not duplicate a completed mutation or provider job."),
            ("AGT-006", "Stream concise task state, decisions, requested approvals, output, verification, and failure reason.", "P0", "The user can understand why a task is waiting or failed without reading raw logs."),
            ("AGT-007", "Package mutations as reviewable changesets with evidence and accept, reject, or partial-apply controls.", "P0", "Rejected work does not remain in the accepted scene or release build."),
            ("AGT-008", "Allow two writing agents on proven-disjoint scopes during private alpha.", "P1", "A two-agent acceptance scenario completes with zero lost edits and a complete audit trail."),
            ("AGT-009", "Use isolated Git worktrees or equivalent isolation for code-writing tasks.", "P1", "Concurrent code changes are independently testable before integration."),
            ("AGT-010", "Keep the initial OpenCode integration behind an internal worker interface.", "P0", "A task does not depend on OpenCode-specific UI or persisted data structures outside the adapter."),
        ),
    )

    add_requirement_section(
        doc,
        "7.4 Asset Foundry",
        "Asset creation must turn a brief into an approved, editable, efficient game asset with traceable lineage rather than merely returning a download link.",
        (
            ("AST-001", "Create a structured asset brief with intended role, style references, dimensions, scale, topology, animation, collision, and target profile.", "P0", "The brief is saved, versioned, and attached to every generated candidate."),
            ("AST-002", "Submit provider-neutral asynchronous jobs and begin with one production provider integration.", "P0", "Jobs survive polling failures and Core restart without duplicate paid generation."),
            ("AST-003", "Generate and compare multiple candidates with preview, provenance, cost, and provider metadata.", "P0", "The user can select, reject, or request variants before import."),
            ("AST-004", "Preserve immutable originals and content-addressed lineage through every processing step.", "P0", "Every production asset can trace back to the source job, brief, and original artifact."),
            ("AST-005", "Process orientation, units, pivot, naming, mesh cleanup, material normalization, texture packaging, and supported compression.", "P0", "The processed GLB imports at the expected scale and orientation with valid materials."),
            ("AST-006", "Add collision, LODs, rig checks, animation checks, and engine-specific import profiles when required by the brief.", "P1", "The asset meets its declared technical profile without manual file surgery."),
            ("AST-007", "Validate geometry, materials, textures, animation, performance budgets, licenses, and parser safety before approval.", "P0", "A failed check blocks production approval and provides a repair reason."),
            ("AST-008", "Require explicit production approval and distinguish placeholder, candidate, approved, deprecated, and rejected states.", "P0", "Only approved assets can enter a release build unless the build is marked development-only."),
            ("AST-009", "Import an approved asset and safely replace placeholder or prior references.", "P0", "The reference replacement preserves stable scene identity and supported overrides."),
            ("AST-010", "Enforce per-task and per-project cost ceilings and retain terms or license reference metadata.", "P1", "A job exceeding its ceiling pauses for approval; exported metadata includes provenance."),
        ),
    )
    base.add_figure(
        doc,
        asset_pipeline,
        "Asset Foundry converts provider output into validated, approved, engine-ready content.",
        "Pipeline diagram from asset brief through generation candidates, processing, validation, human approval, import, and versioned replacement.",
        width=6.35,
    )

    add_requirement_section(
        doc,
        "7.5 Style Packs and immersion",
        "A Style Pack is versioned production data that keeps the world coherent across generation, direct editing, renderer settings, and validation.",
        (
            ("STY-001", "Create, version, duplicate, compare, and apply a Style Pack at project and scene scope.", "P0", "A selected version can be reproduced after restart and in a release build."),
            ("STY-002", "Store palette, material language, shape language, texture density, asset references, lighting intent, and prohibited patterns.", "P0", "Asset briefs and validation can resolve the active style constraints."),
            ("STY-003", "Bind renderer-specific tone mapping, color management, shadows, fog, post-processing, and quality tiers behind the chosen renderer adapter.", "P0", "The reference scene reproduces the approved visual direction on the target profile."),
            ("STY-004", "Coordinate camera behavior, motion language, particles, screen feedback, and interaction feedback guidance.", "P1", "The vertical slice passes the style and responsiveness review checklist."),
            ("STY-005", "Define ambient layers, spatial audio zones, music direction, and event-sound guidance even when audio generation is external.", "P1", "A playtest proves correct zone and event triggering with no missing critical feedback."),
            ("STY-006", "Support explicit object and scene overrides without silently changing the base Style Pack.", "P0", "Overrides are visible, reviewable, and survive Style Pack reapplication."),
            ("STY-007", "Capture reference screenshots and compare material visual changes against approved checkpoints.", "P1", "A material or post-process regression produces review evidence before release."),
        ),
    )

    add_requirement_section(
        doc,
        "7.6 Playtesting, performance, and publishing",
        "Caliber must close the loop from change to evidence. A feature is not complete merely because code or a scene file changed.",
        (
            ("TST-001", "Launch and stop an instrumented playtest without leaving the studio.", "P0", "The selected build reaches an interactive state and reports a clean or failed launch."),
            ("TST-002", "Record or replay supported input sequences and read declared game-state checkpoints.", "P0", "A reference scenario replays deterministically enough to localize and verify a known regression."),
            ("TST-003", "Capture synchronized frames, logs, runtime errors, game state, frame timing, and build identity as test evidence.", "P0", "Every failed reference test creates a reproducible evidence package with a navigable timeline."),
            ("TST-004", "Create a scoped Repair task directly from failed evidence.", "P0", "The repair task contains the failing build, scenario, relevant logs, and acceptance condition."),
            ("PERF-001", "Define a target profile with browser, viewport, reference hardware, renderer path, and content budgets.", "P0", "Every performance report identifies the profile used rather than claiming a universal score."),
            ("PERF-002", "Display frame time, frame pacing, load timing, visible objects, draw calls, triangles, textures, and selected memory indicators.", "P0", "The reference scene exposes budget status in the studio and build report."),
            ("PERF-003", "Track performance deltas by accepted changeset and flag budget regressions.", "P1", "A controlled regression identifies the responsible revision and affected metric."),
            ("PERF-004", "Block release when a required target profile or critical frame-time check fails.", "P0", "Release cannot be marked passing without the required evidence or an explicit waiver."),
            ("PUB-001", "Produce reproducible development and release builds with manifest, version, assets, Style Pack, and source revision.", "P0", "Rebuilding the same accepted state yields the same manifest and equivalent output."),
            ("PUB-002", "Publish a shareable preview URL or a local equivalent during technical alpha.", "P0", "A clean supported browser can load and play the approved vertical slice."),
        ),
    )

    add_requirement_section(
        doc,
        "7.7 Frame evidence and animation diagnosis",
        "Frame review is a first-class production surface. The agent must ground every animation diagnosis in synchronized visual and runtime evidence, then prove the repair by replaying the same sequence.",
        (
            ("ANM-001", "Capture a deterministic frame trace for a selected playtest window, binding every frame to timecode, input, build and revision, game state, active animation state, clip time, root transform, and available bone, socket, or contact transforms.", "P0", "Scrubbing any captured frame recovers its exact build, time, input, object, animation, and state context."),
            ("ANM-002", "Provide frame stepping, variable-speed playback, zoom, loop range, and side-by-side or onion-skin comparison with an accepted reference or prior build.", "P0", "A user isolates and compares the exact failure interval without rerunning the scenario."),
            ("ANM-003", "Detect supported animation failures including clipping or interpenetration, foot sliding, pose popping, jitter, incorrect timing or easing, root-motion drift, contact loss, camera discontinuity, and missing event synchronization.", "P0", "Seeded fixtures identify the correct defect category, affected object or bone, and frame range within the declared tolerance."),
            ("ANM-004", "Require agent findings to cite frame range and timecode, affected object or bone, animation state and clip, visual evidence, relevant telemetry, confidence, and a concise causal explanation.", "P0", "Every diagnosis jumps to cited evidence; an unsupported statement such as 'looks wrong' cannot authorize a mutation."),
            ("ANM-005", "Convert an accepted diagnosis into a bounded Repair task while retaining human approval before production mutation.", "P0", "The task contains the exact reproduction window, expected behavior, resource scope, proposed checks, and source evidence."),
            ("ANM-006", "Replay the same deterministic sequence after repair and compare it with the baseline and accepted reference.", "P0", "Acceptance is blocked until the original defect is absent from its cited interval and required visual, state, and performance regressions pass."),
        ),
    )

    add_requirement_section(
        doc,
        "7.8 Engine adapter contract",
        "The adapter contract allows Caliber to expand without pretending that engines share one universal scene representation.",
        (
            ("ADP-001", "Connect, authenticate where needed, and report project and adapter capabilities.", "P0", "The studio disables unsupported operations and explains the missing capability."),
            ("ADP-002", "Enumerate scenes and resources and inspect selection, identity, revision, dependencies, and semantic role.", "P0", "A task can scope work without parsing opaque native files through chat."),
            ("ADP-003", "Apply an idempotent transactional changeset against expected revisions.", "P0", "Duplicate requests do not duplicate mutations; stale revisions produce a conflict."),
            ("ADP-004", "Use native undo transactions or an explicit compensating reversal.", "P0", "An accepted adapter mutation has a tested reversal path."),
            ("ADP-005", "Import approved assets and preserve stable Caliber identity and native references.", "P0", "Import and replacement pass the adapter contract fixture."),
            ("ADP-006", "Launch playtests, capture evidence, validate, and build through native engine behavior.", "P2", "The Godot adapter passes the shared evidence and build contract."),
            ("ADP-007", "Keep engine-native scenes and resources authoritative for advanced engine features.", "P0", "The shared model stores intent and identity, not a lossy lowest-common-denominator scene."),
        ),
    )

    base.add_heading(doc, "8. Critical user journeys", 1)
    journeys = (
        (
            "8.1 Concept to playable slice",
            (
                "The user creates a project, selects a starter and Style Pack, and enters the desired experience.",
                "Caliber proposes a bounded milestone plan and identifies parallel code, scene, and asset work.",
                "The user approves the initial scopes; the game remains playable while tasks progress.",
                "Caliber integrates verified changes, presents the result, and preserves a reversible revision.",
                "The user plays the loop and converts feedback into the next scoped task.",
            ),
            "A first coherent playable scene exists without manual infrastructure setup, and every accepted change is traceable.",
        ),
        (
            "8.2 Direct edit while agents work",
            (
                "Two agents hold disjoint task scopes while the user selects a scene object.",
                "The user changes transform and material values through gizmos and the inspector.",
                "The direct edit acquires user-priority ownership, commits a revision, and updates the viewport immediately.",
                "An agent touching a stale or overlapping resource rebases or enters Review instead of overwriting.",
                "The final integrated state retains the direct edit and both valid agent outputs.",
            ),
            "The acceptance scenario records zero lost edits, complete changesets, and an understandable conflict outcome if overlap is intentional.",
        ),
        (
            "8.3 Brief to approved game asset",
            (
                "The user opens Asset Foundry from a placeholder and completes or approves the structured brief.",
                "Caliber submits an asynchronous provider job and returns multiple candidates with cost and provenance.",
                "The user selects a candidate; Caliber processes scale, pivot, materials, textures, collision, and budgets.",
                "Automated validation and technical-art review determine whether the candidate can be approved.",
                "The approved asset replaces the placeholder while preserving stable scene identity and overrides.",
            ),
            "The production asset is editable, efficient, licensed or traceable, and reproducible from its lineage record.",
        ),
        (
            "8.4 Failure to verified repair",
            (
                "A build, runtime, playtest, visual, or performance check fails.",
                "Caliber captures the failing revision, scenario, logs, screenshot, state, and profile.",
                "The user or Director creates a Repair task with the smallest relevant scope.",
                "The assigned worker fixes the issue and reruns the failed scenario plus regressions.",
                "The user reviews the changeset and evidence before accepting the repair.",
            ),
            "The repaired build passes the original scenario and does not hide or discard the failed evidence.",
        ),
        (
            "8.5 Animation failure to proven repair",
            (
                "The user records or selects a deterministic playtest where an animation looks wrong.",
                "Caliber captures a synchronized frame trace with rendered frames, input, game state, animation state, clip time, root motion, and available bone or contact telemetry.",
                "The agent marks the exact bad frame range, affected object or bone, defect category, evidence, confidence, and likely cause.",
                "The user inspects the cited frames and approves a bounded Repair task rather than an unconstrained rewrite.",
                "Caliber replays the identical input after repair and compares the same interval against the baseline and accepted reference.",
            ),
            "A seeded animation defect is localized to the correct interval and component, repaired, and proven absent without introducing a visual, state, or performance regression.",
        ),
        (
            "8.6 Validate and publish",
            (
                "The user selects a release candidate and required target profile.",
                "Caliber runs build, clean-load, smoke-playtest, runtime-error, asset-state, and performance checks.",
                "Blocking failures return to scoped repair; approved waivers are explicit and recorded.",
                "The release manifest pins project revision, assets, Style Pack, renderer, profile, and build settings.",
                "Caliber publishes a preview URL and retains the report and artifact identity.",
            ),
            "A clean supported browser loads a reproducible playable build whose required release gates pass.",
        ),
    )
    for title, steps, success in journeys:
        base.add_heading(doc, title, 2)
        journey_num_id = base.add_numbering_definition(doc, "decimal")
        for step in steps:
            base.add_numbered(doc, step, journey_num_id)
        base.add_callout(doc, "Success condition", success, fill=base.LIGHT, accent=base.BLUE)

    base.add_heading(doc, "9. Product rules and domain model", 1, new_page=True)
    base.add_heading(doc, "Authoritative records", 2)
    domain_rows = (
        ("Project", "Product intent, target profiles, engine bindings, active Style Pack, permissions, cost limits, and release history."),
        ("Resource", "Stable Caliber ID, native locator, type, semantic role, revision, dependencies, active lease, and validation status."),
        ("Task", "Objective, owner, scope, dependencies, state, decisions, constraints, cost, output, and evidence."),
        ("Changeset", "Expected revisions, mutations, author, task, evidence, approval, integration status, and reversal path."),
        ("Asset", "Brief, immutable original, candidates, processor lineage, metadata, validation, approval, versions, and references."),
        ("Style Pack", "Versioned visual and immersion intent plus renderer, asset, camera, feedback, and audio bindings."),
        ("Playtest", "Build, scenario, inputs, state assertions, screenshots, logs, performance profile, result, and repair link."),
        ("Frame trace", "A deterministic, time-indexed evidence record joining rendered frames to inputs, build identity, game state, animation state, clip time, transforms, contacts, performance, annotations, and repair links."),
        ("Build", "Source revision, manifest, renderer or engine, assets, Style Pack, profile, validation report, artifact, and publication."),
    )
    base.add_table(doc, ("Record", "Required meaning"), domain_rows, (1700, 7660))
    base.add_heading(doc, "State rules", 2)
    for item in (
        "Accepted state changes are append-only events plus derived current state; important history is not reconstructed from chat.",
        "Every mutation carries an idempotency key and expected revision set.",
        "Optimistic concurrency is allowed for structured text and disjoint resources; fragile scenes and binary assets may require exclusive ownership.",
        "A user edit outranks an agent lease. Integration never resolves a conflict through silent last-write-wins behavior.",
        "Original provider output is immutable. Processing creates derived versions; production approval selects a version without destroying lineage.",
        "Development builds may include placeholders; release builds may not include unapproved assets unless an explicit waiver is recorded.",
        "A visual diagnosis without cited frame evidence cannot authorize a production mutation; the same deterministic replay must verify an accepted repair.",
        "Chat is context, not authority. Tasks, changesets, assets, style versions, evidence, and approvals are authoritative product records.",
    ):
        base.add_bullet(doc, item, bullet_num_id)

    base.add_heading(doc, "10. Non-functional requirements", 1, new_page=True)
    base.add_heading(doc, "Performance and responsiveness", 2)
    nfr_rows = (
        ("NFR-P01", "Viewport target", "60 FPS at 1080p on the defined reference hardware and scene; 30 FPS is the hard development floor."),
        ("NFR-P02", "Interaction", "Selection feedback within 100 ms and local inspector acknowledgement within 250 ms on the reference project."),
        ("NFR-P03", "Project readiness", "Starter project reaches an interactive viewport within 10 seconds at P50 and 30 seconds at P95 after warm installation."),
        ("NFR-P04", "Progressive loading", "Large optional content does not block the first playable frame and exposes visible loading state."),
        ("NFR-P05", "Evidence", "Every performance result records browser, viewport, reference hardware, build, renderer path, and content profile."),
        ("NFR-P06", "Frame capture", "Frame-trace capture declares its timing overhead, preserves source frame numbers and timecodes, and loads long traces progressively without blocking the live viewport."),
    )
    base.add_table(doc, ("ID", "Area", "Initial target"), nfr_rows, (1200, 2100, 6060))
    base.add_heading(doc, "Reliability and recovery", 2)
    reliability_rows = (
        ("NFR-R01", "No accepted edit is lost after process restart or worker crash."),
        ("NFR-R02", "Duplicate task, changeset, build, or provider-job requests are idempotent."),
        ("NFR-R03", "A corrupted or unavailable subsystem degrades explicitly and does not corrupt the last valid project state."),
        ("NFR-R04", "Core event replay reconstructs accepted current state in automated recovery tests."),
        ("NFR-R05", "The private-alpha target is at least 99% crash-free studio sessions after the stabilization milestone."),
    )
    base.add_table(doc, ("ID", "Requirement"), reliability_rows, (1500, 7860))
    base.add_heading(doc, "Compatibility and portability", 2)
    for item in (
        "Support a declared desktop-browser baseline; prefer WebGPU with a tested WebGL 2 fallback where the selected renderer and project profile allow it.",
        "Keep renderer-specific details behind a Web renderer adapter and choose one production renderer before alpha.",
        "Use glTF or GLB as the initial runtime interchange format for supported 3D assets, with processing and validation before import.",
        "Keep Core protocols versioned so browser Studio and future native-engine plugins can negotiate capabilities.",
        "Do not require a desktop wrapper for the Web technical alpha; add a native hub only for proven product needs.",
    ):
        base.add_bullet(doc, item, bullet_num_id)
    base.add_heading(doc, "Accessibility and usability", 2)
    for item in (
        "Keyboard-reachable primary navigation, prompts, tasks, approvals, inspector fields, play controls, and build actions.",
        "Visible focus, readable contrast, non-color-only states, text zoom support, and meaningful error messages.",
        "A hierarchy and inspector alternative for essential viewport actions that otherwise depend on precise pointer input.",
        "Status language understandable without reading raw model or tool traces.",
    ):
        base.add_bullet(doc, item, bullet_num_id)

    base.add_heading(doc, "11. Trust, security, and privacy", 1)
    base.add_heading(doc, "Trust model", 2)
    base.add_body(
        doc,
        "Caliber earns trust by making authority, scope, cost, mutation, evidence, and recovery visible. The user should never need to guess whether a model can access secrets, mutate an unrelated file, pay a provider, or overwrite a scene."
    )
    security_rows = (
        ("SEC-001", "Local Core binds to loopback by default and clients use short-lived authentication."),
        ("SEC-002", "Provider and model credentials remain in Core or the operating-system keychain and never enter the game preview."),
        ("SEC-003", "Game previews run in a sandbox without unrestricted host filesystem, shell, or credential access."),
        ("SEC-004", "Agent tools are allowlisted and scoped by project, task, resource, network, time, and cost."),
        ("SEC-005", "High-risk actions such as broad deletion, external publishing, dependency replacement, or cost increases require approval."),
        ("SEC-006", "Imported assets receive signature, type, size, path, archive, and parser-safety checks."),
        ("SEC-007", "Projects declare whether external providers, specific vendors, telemetry, or cloud storage are permitted."),
        ("SEC-008", "Logs and analytics exclude source, prompts, credentials, and private assets by default unless the user explicitly opts in."),
        ("SEC-009", "Every mutation, approval, provider job, build, and publication action has an auditable actor and timestamp."),
    )
    base.add_table(doc, ("ID", "Requirement"), security_rows, (1500, 7860))

    base.add_heading(doc, "12. Analytics and success metrics", 1)
    base.add_callout(
        doc,
        "North-star metric",
        "Weekly number of user-approved, playtested, performance-passing game improvements.",
        fill=base.PALE_TEAL,
        accent=base.TEAL,
    )
    base.add_heading(doc, "Activation and value", 2)
    metric_rows = (
        ("Time to first playable", "Median time from new project to interactive starter and from first prompt to visible accepted change.", "Starter <= 10 minutes; first accepted prompt change <= 5 minutes."),
        ("Core-loop completion", "Users who prompt or edit, play, review evidence, and accept a change in the first session.", "At least 70% of supported private-alpha sessions."),
        ("Direct-edit durability", "Accepted direct edits that survive save, reload, play transitions, agent integration, and restart.", "100% in acceptance fixtures; zero known silent-loss defects."),
        ("Task success", "Scoped tasks accepted without manual code repair after required verification.", "At least 80% for curated alpha workflows."),
        ("Asset approval", "Briefs producing an approved import within three candidate or repair rounds.", "At least 60% for supported asset classes."),
        ("Animation diagnosis", "Seeded defects localized to the correct frame interval, affected component, and supported defect category, then cleared by the same replay.", "At least 80% of curated fixtures; zero accepted diagnoses without cited frames."),
        ("Release success", "Approved release candidates that pass clean-load, smoke-play, runtime, asset, and performance gates.", "At least 90% after one scoped repair cycle."),
    )
    base.add_table(doc, ("Metric", "Definition", "Initial alpha target"), metric_rows, (2200, 4300, 2860))
    base.add_heading(doc, "Guardrail metrics", 2)
    for item in (
        "Lost-edit count, conflict rate, rebase rate, duplicate-operation rate, recovery success, and blocked time.",
        "Agent correction rate, approval rejection rate, unverified change count, and repair recurrence.",
        "Animation localization error in frames, false-positive rate, time to grounded diagnosis, repair recurrence, and frame-capture overhead.",
        "Asset cost per approved import, provider failure rate, license-metadata completeness, and technical-art repair rate.",
        "Frame-time pass rate, clean-load pass rate, runtime errors, budget violations, and visual regression rate.",
        "Permission denials, out-of-scope attempts, credential exposure incidents, and unintended publication events.",
    ):
        base.add_bullet(doc, item, bullet_num_id)
    base.add_heading(doc, "Telemetry principles", 2)
    base.add_body(
        doc,
        "Private alpha begins local-first. Product analytics use explicit opt-in and event schemas that avoid prompts, code, private assets, and credentials. The product remains usable when analytics are disabled."
    )

    base.add_heading(doc, "13. Delivery and rollout plan", 1, new_page=True)
    roadmap_rows = (
        ("0. Architecture and quality spike", "2 weeks", "Renderer bakeoff, Core event spike, direct transform, frame-trace feasibility, Style Pack v0, provider sample", "Renderer, frame evidence, and provider decision evidence"),
        ("1. Web editor foundation", "4 weeks", "Projects, shell, viewport, hierarchy, selection, gizmos, inspector, modes, undo, persistence", "Small scene built without code"),
        ("2. Safe AI loop", "4 weeks", "OpenCode worker, tasks, scopes, leases, revisions, changesets, review, recovery", "Human and agent edit safely"),
        ("3. Asset Foundry alpha", "5 weeks", "Briefs, provider jobs, candidates, lineage, processing, validation, approval, replacement", "Approved asset replaces placeholder"),
        ("4. Immersion and release", "5 weeks", "Style binding, camera, feedback, audio zones, frame timeline, animation diagnosis, repair replay, performance, build, publish", "New user completes reference flow"),
        ("5. Private-alpha stabilization", "2-4 weeks", "Onboarding, recovery UI, two-agent fixture, accessibility, crash and performance fixes", "External cohort launch gate passes"),
        ("6. Godot adapter", "6-8 weeks", "Native plugin, selection, undo, GLB import, playtest, validation, build", "Shared adapter contract passes natively"),
    )
    base.add_table(doc, ("Phase", "Duration", "Primary deliverable", "Exit gate"), roadmap_rows, (2300, 1200, 3460, 2400))
    base.add_heading(doc, "Team assumption", 2)
    for item in (
        "Product and AI systems lead, plus a senior Web3D and graphics engineer.",
        "Rust and infrastructure engineer.",
        "Technical artist or tools artist, at least part time from the first asset milestone.",
        "Product designer with strong editor, workflow, and accessibility experience, full time or embedded by phase 1.",
    ):
        base.add_bullet(doc, item, bullet_num_id)
    base.add_body(
        doc,
        "A focused team can target the private-alpha gate in approximately 22-24 calendar weeks including stabilization. A solo implementation should plan for substantially longer and reduce parallel scope rather than skipping safety or asset-quality gates."
    )
    base.add_heading(doc, "Rollout sequence", 2)
    rollout_num_id = base.add_numbering_definition(doc, "decimal")
    for item in (
        "Internal dogfood on one golden vertical slice and one asset benchmark set.",
        "Design partners: three to five technical creators working from the supported starter and target profile.",
        "Private alpha: ten to twenty users after P0 data-loss, security, recovery, and performance gates pass.",
        "Godot design partners only after the Web adapter contract is stable and the native plugin passes contract fixtures.",
    ):
        base.add_numbered(doc, item, rollout_num_id)

    base.add_heading(doc, "14. Release acceptance criteria", 1, new_page=True)
    base.add_heading(doc, "Web technical alpha", 2)
    acceptance_rows = (
        ("Core experience", "User creates or opens a project, prompts a change, directly edits the scene, plays, reviews, and restores after restart."),
        ("Direct editing", "Selection, transforms, supported inspector edits, undo, save, reload, and identity persistence pass the golden project fixture."),
        ("Agent safety", "A user and at least one writing agent edit safely; stale revision, duplicate request, cancel, crash, and recovery fixtures pass."),
        ("Asset Foundry", "One supported 3D asset goes from brief to candidates, processing, validation, approval, import, and reference replacement with lineage."),
        ("Style and immersion", "One versioned Style Pack controls the approved reference scene and required visual, camera, feedback, and audio-zone checks pass."),
        ("Playtest", "A reference scenario captures inputs, screenshots, logs, state, and frame evidence; a known failure creates a reproducible repair task."),
        ("Frame diagnosis", "A deterministic scenario captures synchronized frames and animation telemetry; a seeded defect is localized to the exact interval and component; the same replay proves the approved repair."),
        ("Performance", "The vertical slice passes the declared 1080p target profile and exposes useful budget evidence."),
        ("Publishing", "A reproducible release build loads and plays in a clean supported browser from a preview URL or technical-alpha equivalent."),
        ("Trust", "No open P0 data-loss, credential-exposure, silent-overwrite, scope-escape, or unintended-publication defect remains."),
    )
    base.add_table(doc, ("Gate", "Acceptance criterion"), acceptance_rows, (2100, 7260))
    base.add_heading(doc, "Private alpha", 2)
    for item in (
        "Two disjoint writing agents pass the concurrency fixture while direct user edits retain priority.",
        "Supported asset classes meet the approval, import, provenance, cost, and technical-art quality thresholds.",
        "Supported animation defect classes meet the localization, false-positive, capture-overhead, and replay-verified repair targets.",
        "Critical onboarding, recovery, approval, and release flows pass usability and accessibility review.",
        "Crash-free session, build success, task success, and performance pass rates meet the initial targets for two consecutive dogfood weeks.",
        "Known limitations, supported browser and hardware profile, privacy behavior, provider costs, and project backup guidance are documented in-product.",
    ):
        base.add_bullet(doc, item, bullet_num_id)

    base.add_heading(doc, "15. Dependencies, risks, and mitigations", 1, new_page=True)
    base.add_heading(doc, "Critical dependencies", 2)
    dependency_rows = (
        ("Web renderer", "Three.js or Babylon.js selected by benchmark; supported browser APIs and WebGPU/WebGL behavior."),
        ("Agent runtime", "OpenCode SDK/server behavior, configured model providers, and a stable internal worker adapter."),
        ("Asset providers", "At least one reliable API, usage rights, rate limits, job status, output format, and predictable cost controls."),
        ("Technical-art pipeline", "Blender command-line processing and validated glTF/GLB tooling for production-ready assets."),
        ("Local Core", "Rust service, SQLite event and state persistence, artifact storage, authentication, tasks, leases, and protocol versioning."),
        ("Frame evidence pipeline", "Deterministic input replay, renderer capture, animation and transform telemetry, artifact storage, comparison tools, and a vision-analysis boundary."),
        ("Publishing", "A reproducible static Web build and preview-hosting path that does not expose provider or project secrets."),
    )
    base.add_table(doc, ("Dependency", "Required capability"), dependency_rows, (2200, 7160))
    base.add_heading(doc, "Major risks", 2)
    risk_rows = (
        ("Four engines too early", "No platform becomes reliable", "Ship Web, prove the adapter contract, then Godot; gate Unity and Unreal on demand and capacity."),
        ("Renderer indecision", "Dual maintenance delays alpha", "Run the fixed ten-day bakeoff, choose one production renderer, and defer the second."),
        ("Agent overwrites user", "Trust is lost", "User-priority leases, expected revisions, transactions, conflict UI, evidence, and acceptance fixtures."),
        ("Agent invents visual diagnosis", "Wrong repairs damage motion and trust", "Frame-cited findings, synchronized telemetry, seeded fixtures, confidence, human approval, and same-sequence replay proof."),
        ("Generated assets look random", "Game lacks coherent art direction", "Style Pack, references, candidate comparison, technical-art review, and explicit approval."),
        ("Provider output is weak", "Bad deformation or performance", "Processing, validation, repair loops, provider routing, and supported asset classes."),
        ("Beautiful but slow Web scene", "Product breaks its promise", "Target profiles, budgets, deltas, LODs, compression, culling, streaming, and release blocks."),
        ("Universal scene abstraction", "Native engine features are lost", "Share intent and identity while keeping detailed scenes engine-native."),
        ("Security boundary is weak", "Credentials or files are exposed", "Local Core, preview sandbox, scoped tools, keychain, import safety, audit, and approvals."),
        ("UX becomes an IDE cockpit", "Creators cannot build momentum", "Progressive disclosure, product-language status, one golden flow, and first-run usability tests."),
        ("Costs are unpredictable", "Users avoid asset or agent workflows", "Visible estimates, ceilings, approval thresholds, caching, idempotency, and per-output metrics."),
    )
    base.add_table(doc, ("Risk", "Consequence", "Mitigation"), risk_rows, (2600, 2460, 4300))

    base.add_heading(doc, "16. Open decisions and decision gates", 1, new_page=True)
    decision_rows = (
        ("D1", "Production Web renderer", "End of day 10", "Visual quality, editor ergonomics, loading, fallback, performance, bundle, systems, and team productivity.", "Open"),
        ("D2", "Default 3D provider by asset class", "End of phase 0", "Same briefs and references tested across Tripo and Meshy; compare approval quality, cost, latency, editability, and terms.", "Open"),
        ("D3", "Technical-alpha publishing path", "End of phase 1", "Local preview vs hosted URL, privacy, build reproducibility, cost, and feedback workflow.", "Open"),
        ("D4", "Native desktop hub", "After private-alpha evidence", "Add Tauri only if process control, keychain, deep links, distribution, offline use, or performance justify it.", "Deferred"),
        ("D5", "Godot start gate", "After Web private alpha", "Core, concurrency, Asset Foundry, Style Pack, evidence schema, and adapter contract must be stable.", "Gated"),
        ("D6", "Unity and Unreal sequence", "After Godot adapter proof", "Require user demand, native-tool expertise, binary-safe source control, maintenance capacity, and commercial rationale.", "Gated"),
        ("D7", "Primary market positioning", "Before public beta", "Technical creator vs designer-led team; determine default code visibility, onboarding, pricing, and support expectations.", "Open"),
        ("D8", "Imported-project support", "Before public beta", "Evaluate risk and value of arbitrary repositories versus Caliber-created project profiles and migration tooling.", "Deferred"),
        ("D9", "Frame capture and vision-analysis architecture", "End of phase 0", "Compare deterministic capture fidelity, timing overhead, artifact size, animation telemetry coverage, browser constraints, privacy, cost, and local versus provider vision analysis.", "Open"),
    )
    base.add_table(doc, ("ID", "Decision", "Deadline", "Evidence required", "Status"), decision_rows, (650, 1900, 1400, 4310, 1100))
    base.add_heading(doc, "Decision rule", 2)
    base.add_body(
        doc,
        "A gate closes with written evidence, a named owner, the selected option, rejected alternatives, consequences, and a review date. An unresolved gate cannot silently turn into permanent dual support."
    )

    base.add_heading(doc, "Appendix A. Requirement index by release", 1, new_page=True)
    index_rows = (
        ("P0", "PRJ-001-005; UI-001, UI-003, UI-005; EDT-001-007; AGT-001-007, AGT-010; AST-001-005, AST-007-009; STY-001-003, STY-006; TST-001-004; ANM-001-006; PERF-001-002, PERF-004; PUB-001-002; ADP-001-005, ADP-007"),
        ("P1", "UI-002, UI-004; EDT-008; AGT-008-009; AST-006, AST-010; STY-004-005, STY-007; PERF-003"),
        ("P2", "ADP-006 and later native-engine, cloud-team, provider, marketplace, and console expansion approved through decision gates"),
    )
    base.add_table(doc, ("Release", "Requirement IDs and scope"), index_rows, (1400, 7960))
    base.add_heading(doc, "Traceability rule", 2)
    base.add_body(
        doc,
        "Every implementation task references one or more requirement IDs. Every P0 requirement has an acceptance fixture or product review method, and every release gate links to the evidence created by those fixtures."
    )

    base.add_heading(doc, "Appendix B. Product glossary", 1)
    glossary_rows = (
        ("Adapter", "A renderer- or engine-native integration that exposes Caliber capabilities while preserving native resources and workflows."),
        ("Asset Foundry", "The provider-neutral workflow for briefs, generation, candidates, processing, validation, approval, import, lineage, and replacement."),
        ("Caliber Core", "The durable local service for projects, events, tasks, permissions, revisions, leases, changesets, assets, evidence, and builds."),
        ("Changeset", "A scoped, idempotent set of mutations with expected revisions, evidence, approval state, and reversal path."),
        ("Evidence", "Reproducible output used to evaluate work: tests, screenshots, logs, state, performance, visual checks, and build reports."),
        ("Frame trace", "Synchronized frame-by-frame evidence that ties what was rendered to the exact input, build, runtime state, animation state, transforms, contacts, timing, and annotations."),
        ("Lease", "Temporary ownership of mutable resources that prevents conflicting writers and gives active human editing priority."),
        ("MCP tool", "A bounded agent-facing operation for project, scene, asset, test, or build work; it is not an engine or scene format."),
        ("Style Pack", "Versioned visual and immersion intent plus bindings for assets, materials, renderer, camera, effects, feedback, and audio."),
        ("Target profile", "The browser, viewport, reference hardware, renderer path, scene class, and quality or performance budgets used for evidence."),
        ("Vertical slice", "A short, polished, representative game experience proving the complete gameplay and production loop at controlled scope."),
    )
    base.add_table(doc, ("Term", "Definition"), glossary_rows, (2000, 7360))
    base.add_callout(
        doc,
        "Final product statement",
        "Caliber Studio is the production system that lets creators combine direct visual authorship with trustworthy asynchronous AI work, frame-grounded diagnosis, production-ready content, coherent style, engine-native execution, and verified shipping.",
        fill=base.PALE_TEAL,
        accent=base.TEAL,
    )

    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.core_properties.title = "Caliber Studio Product Requirements Document"
    doc.core_properties.subject = "Web-first AI-native game production workspace"
    doc.core_properties.author = "Caliber Studio"
    doc.core_properties.keywords = (
        "Caliber Studio, PRD, AI gaming, Web3D, direct editing, multi-agent, Asset Foundry, Style Pack, Godot"
    )
    doc.core_properties.comments = "Founder-ready product requirements document version 1.1."

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
